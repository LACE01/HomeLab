"""Combined audit-ready export for a Risk Register entry -- pulls together the
risk itself plus everything cross-linked to it (findings, assets, Albert
network-monitoring alerts, and exceptions/risk-acceptances) into one Word
document. This is the natural "give this to the auditor" artifact once a risk
has accumulated links from the various places that can attach to it (Albert
alert triage, exception requests, manual linking on Risk Detail).

Reuses the same python-docx patterns as incident_response.build_case_docx
(heading levels, Light Grid Accent 1 tables, a BytesIO stream returned for a
StreamingResponse) rather than sharing code with it directly, since the two
reports have fairly different shapes (case narrative vs. a register entry with
several distinct linked-item tables).
"""
import io
from datetime import datetime, timezone


def _fmt_dt(value):
    if not value:
        return "-"
    try:
        return str(value)[:10]
    except Exception:
        return str(value)


def build_risk_export_docx(risk: dict, findings: list, assets: list, albert_alerts: list, exceptions: list, ir_cases: list = None) -> io.BytesIO:
    ir_cases = ir_cases or []
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading(f"Risk Register Report — {risk.get('title', '')}", level=0)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Risk ID: {risk.get('id', '')}    Category: {risk.get('category', '-')}    "
        f"Status: {risk.get('status', '-')}    Owner: {risk.get('owner') or 'Unassigned'}"
    ).italic = True
    generated = doc.add_paragraph()
    generated.add_run(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}").font.size = Pt(9)

    doc.add_heading("Description", level=1)
    doc.add_paragraph(risk.get("description") or "-")

    doc.add_heading("Risk Scoring", level=1)
    score_table = doc.add_table(rows=1, cols=4)
    score_table.style = "Light Grid Accent 1"
    hdr = score_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "", "Likelihood", "Impact", "Score / Band"
    row = score_table.add_row().cells
    row[0].text = "Inherent"
    row[1].text = str(risk.get("likelihood", "-"))
    row[2].text = str(risk.get("impact", "-"))
    row[3].text = f"{risk.get('inherent_score', '-')} ({risk.get('inherent_band', '-')})"
    if risk.get("residual_score"):
        row2 = score_table.add_row().cells
        row2[0].text = "Residual"
        row2[1].text = str(risk.get("residual_likelihood", "-"))
        row2[2].text = str(risk.get("residual_impact", "-"))
        row2[3].text = f"{risk.get('residual_score', '-')} ({risk.get('residual_band', '-')})"

    doc.add_heading("Treatment", level=1)
    doc.add_paragraph(f"Strategy: {risk.get('treatment_strategy', '-')}")
    doc.add_paragraph(risk.get("treatment_plan") or "-")
    if risk.get("external_reference"):
        doc.add_paragraph(f"External reference (device / domain / website): {risk['external_reference']}")

    doc.add_heading(f"Linked Findings ({len(findings)})", level=1)
    if findings:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text = "Title", "CVE", "Severity", "Status"
        for f in findings:
            c = t.add_row().cells
            c[0].text = f.get("title") or f.get("id", "")
            c[1].text = f.get("cve") or "-"
            c[2].text = f.get("severity") or "-"
            c[3].text = f.get("status") or "-"
    else:
        doc.add_paragraph("No findings linked.")

    doc.add_heading(f"Linked Assets ({len(assets)})", level=1)
    if assets:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text = "Hostname", "IP", "Criticality", "Environment"
        for a in assets:
            c = t.add_row().cells
            c[0].text = a.get("hostname") or a.get("id", "")
            c[1].text = a.get("ip") or "-"
            c[2].text = a.get("criticality") or "-"
            c[3].text = a.get("environment") or "-"
    else:
        doc.add_paragraph("No assets linked.")

    doc.add_heading(f"Linked Albert Network Monitoring Alerts ({len(albert_alerts)})", level=1)
    if albert_alerts:
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = "Time (GMT)", "Alert", "Severity", "Source IP", "Destination IP"
        for al in albert_alerts:
            c = t.add_row().cells
            c[0].text = _fmt_dt(al.get("time_gmt"))
            c[1].text = al.get("alert_message") or "-"
            c[2].text = al.get("severity") or "-"
            c[3].text = al.get("source_ip") or "-"
            c[4].text = al.get("destination_ip") or "-"
    else:
        doc.add_paragraph("No Albert alerts linked.")

    doc.add_heading(f"Linked Exceptions / Risk Acceptances ({len(exceptions)})", level=1)
    if exceptions:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text = "Finding / Target", "Status", "Expires", "Business Justification"
        for e in exceptions:
            c = t.add_row().cells
            c[0].text = e.get("finding_title") or e.get("target_value") or e.get("id", "")
            c[1].text = e.get("status") or "-"
            c[2].text = _fmt_dt(e.get("expires_at"))
            c[3].text = e.get("business_justification") or "-"
    else:
        doc.add_paragraph("No exceptions linked.")

    doc.add_heading(f"Linked Incident Response Cases ({len(ir_cases)})", level=1)
    if ir_cases:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text = "Case #", "Title", "Classification", "Status"
        for c_ in ir_cases:
            c = t.add_row().cells
            c[0].text = c_.get("case_number") or c_.get("id", "")
            c[1].text = c_.get("title") or "-"
            c[2].text = c_.get("classification") or "-"
            c[3].text = c_.get("status") or "-"
    else:
        doc.add_paragraph("No IR cases linked.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
