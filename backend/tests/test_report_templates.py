"""Report layout templates: the report is configuration, not code. Sections can
be reordered, renamed, hidden, configured and added, and every renderer (print
view, shared copy, Word export) reads the same resolved layout."""
import os, sys, asyncio, io
os.environ["MONGO_URL"] = "mongodb://localhost:27017"
os.environ["DB_NAME"] = "test_report_templates"
os.environ["JWT_SECRET"] = "testsecret"
sys.path.insert(0, ".")

from mongomock_motor import AsyncMongoMockClient
import db as db_module
db_module.client = AsyncMongoMockClient()
db_module.db = db_module.client["test_report_templates"]

import server
import auth_utils
from routes import security_reviews as sr_route
sr_route.db = db_module.db
import report_templates as rt
from security_review_docx import build_review_docx

from fastapi.testclient import TestClient
from docx import Document

admin_user = {"id": "u1", "email": "admin@x.com", "role": "admin", "name": "Admin", "teams": []}
app = server.app
app.dependency_overrides[auth_utils.get_current_user] = lambda: admin_user
app.dependency_overrides[auth_utils.get_current_user_optional] = lambda: admin_user
client = TestClient(app)
db = db_module.db


def run(coro):
    return asyncio.get_event_loop().run_until_complete(coro)


# ============ default layout ============

layout = rt.default_layout()
types = [b["type"] for b in layout]
assert types[0] == "header"
# the ask: technical detail right after the executive part, questionnaire LAST
qi = types.index("questionnaire")
assert qi == len(types) - 1, "the questionnaire appendix must be the last section"
assert types.index("linked_assets") < qi and types.index("external_checks") < qi
assert types.index("executive_summary") < types.index("linked_assets"), \
    "executive material comes before the technical detail"
assert all(b.get("title") for b in layout if b["type"] not in ("page_break",))
print("PASS: the stock layout puts executive material first, technical detail next, and the questionnaire appendix LAST")

# ============ validation ============

try:
    rt.validate_blocks([{"type": "not_a_real_block"}])
    raise AssertionError("expected rejection")
except ValueError as e:
    assert "unknown block type" in str(e)
try:
    rt.validate_blocks([{"type": "key_findings"}])
    raise AssertionError("expected rejection")
except ValueError as e:
    assert "header block is required" in str(e)
try:
    rt.validate_blocks([])
    raise AssertionError("expected rejection")
except ValueError:
    pass
ok = rt.validate_blocks([{"type": "header"}, {"type": "key_findings", "options": {"limit": 3}}])
assert ok[1]["options"]["limit"] == 3
assert ok[1]["options"]["show_recommendations"] is True, "defaults fill in around the override"
assert ok[0]["visible"] is True and ok[0]["id"]
print("PASS: layouts reject unknown blocks and a missing header, and normalize options around defaults")

# ============ internal-only blocks can never leak ============

blocks = rt.validate_blocks([
    {"type": "header"}, {"type": "notes"}, {"type": "audit_trail"}, {"type": "key_findings"}])
internal = rt.resolve_layout({"blocks": blocks}, shared=False)
external = rt.resolve_layout({"blocks": blocks}, shared=True)
assert [b["type"] for b in internal] == ["header", "notes", "audit_trail", "key_findings"]
assert "notes" not in [b["type"] for b in external]
assert "audit_trail" not in [b["type"] for b in external]
print("PASS: internal-only blocks (working notes, audit trail) are stripped from the shared copy centrally — "
      "no layout edit can leak them")

hidden = rt.resolve_layout({"blocks": [
    {**blocks[0]}, {**blocks[3], "visible": False}]}, shared=False)
assert [b["type"] for b in hidden] == ["header"]
print("PASS: hiding a block removes it from every renderer")

# consecutive/trailing page breaks collapse so an edit can't produce blank pages
messy = rt.validate_blocks([
    {"type": "page_break"}, {"type": "header"}, {"type": "page_break"},
    {"type": "page_break"}, {"type": "key_findings"}, {"type": "page_break"}])
cleaned = [b["type"] for b in rt.resolve_layout({"blocks": messy}, shared=False)]
assert cleaned == ["header", "page_break", "key_findings"], cleaned
print("PASS: leading, duplicate and trailing page breaks collapse, so an edited layout can't emit blank pages")

# ============ API: versioned saves ============

r = client.get("/api/v1/report-templates/blocks")
cat = r.json()
assert len(cat["blocks"]) >= 15
assert any(b["internal_only"] for b in cat["blocks"])
assert any(b["type"] == "section_heading" for b in cat["blocks"])
assert cat["default_layout"]
print("PASS: the block catalog exposes every renderable section with its description, options and internal-only flag")

r = client.get("/api/v1/report-templates")
assert r.json()["active"]["version"] == 1

custom = [
    {"type": "header", "title": "Eagle County Security Assessment"},
    {"type": "section_heading", "title": "For leadership", "options": {"subtitle": "Read this part."}},
    {"type": "risk_verdict"},
    {"type": "executive_summary", "title": "The short version"},
    {"type": "key_findings", "title": "What we found", "options": {"limit": 3}},
    {"type": "page_break"},
    {"type": "section_heading", "title": "For the technical team"},
    {"type": "linked_assets"},
    {"type": "external_checks", "options": {"panels": "technical"}},
    {"type": "notes"},
    {"type": "questionnaire", "title": "Appendix A — full questionnaire"},
]
r = client.post("/api/v1/report-templates", json={"name": "Eagle County layout", "blocks": custom})
assert r.status_code == 200
assert r.json()["version"] == 2, "saving creates a NEW version rather than editing in place"
saved = r.json()
assert saved["blocks"][0]["title"] == "Eagle County Security Assessment"

r = client.post("/api/v1/report-templates", json={"name": "bad", "blocks": [{"type": "nope"}]})
assert r.status_code == 400
print("PASS: saving a layout creates a new version (old reports stay explainable) and rejects invalid blocks")

# v1 is untouched
v1 = run(db.report_templates.find_one({"key": rt.DEFAULT_KEY, "version": 1}, {"_id": 0}))
assert v1["blocks"][0]["title"] == "Security Review Report"
print("PASS: earlier versions are left intact")

# ============ the saved layout actually drives the report ============

client.get("/api/v1/security-reviews/meta")
review = client.post("/api/v1/security-reviews", json={
    "title": "Acme SaaS", "review_type": "New software purchase (SaaS / COTS / on-prem)",
    "entity_name": "Acme", "entity_domain": "acme.com"}).json()
rid = review["id"]
client.put(f"/api/v1/security-reviews/{rid}/risk-score", json={
    "inherent_likelihood": 4,
    "inherent_impacts": {"confidentiality": 5, "integrity": 3, "availability": 2,
                          "compliance_legal": 4, "reputational": 3},
    "compensating_controls": "MFA everywhere",
    "residual_likelihood": 2,
    "residual_impacts": {"confidentiality": 3, "integrity": 2, "availability": 2,
                          "compliance_legal": 2, "reputational": 2}})
client.post(f"/api/v1/security-reviews/{rid}/findings", json={"description": "No MFA", "severity": "High"})
client.post(f"/api/v1/security-reviews/{rid}/notes", json={"text": "Internal note", "html": "<b>Internal</b> note"})
_app = client.get(f"/api/v1/security-reviews/{rid}").json()["applicable_questions"]
client.put(f"/api/v1/security-reviews/{rid}/responses", json={
    "question_order": _app[0]["order"], "answer": "no", "evidence_text": "vendor confirmed"})

rep = client.get(f"/api/v1/security-reviews/{rid}/report-data").json()
rep_types = [b["type"] for b in rep["layout"]]
assert rep_types[0] == "header"
assert rep["layout"][0]["title"] == "Eagle County Security Assessment", "the report renders under the SAVED layout"
assert rep_types[-1] == "questionnaire"
assert "notes" in rep_types
assert rep["template"]["version"] == 2 and "blocks" not in rep["template"]
kf = next(b for b in rep["layout"] if b["type"] == "key_findings")
assert kf["title"] == "What we found" and kf["options"]["limit"] == 3
print("PASS: report data returns the resolved layout from the saved template, including custom titles and options")

# shared copy resolves the same layout minus internal-only blocks
r = client.post(f"/api/v1/security-reviews/{rid}/share", json={"email": "ext@partner.com"})
token = r.json()["token"]
grant = run(db.security_review_share_grants.find_one({"token": token}, {"_id": 0}))
shared = client.post(f"/api/v1/shared/security-review/{token}/verify", json={"code": grant["code"]}).json()
shared_types = [b["type"] for b in shared["layout"]]
assert "notes" not in shared_types, "the shared layout must never include working notes"
assert "header" in shared_types and "questionnaire" in shared_types
assert shared["notes"] == []
print("PASS: the shared report resolves the same layout with internal-only sections removed")

# ============ Word export follows the layout ============

r = client.get(f"/api/v1/security-reviews/{rid}/export.docx")
assert r.status_code == 200
doc = Document(io.BytesIO(r.content))
heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
assert heads[0] == "Eagle County Security Assessment", heads
assert "For leadership" in heads and "For the technical team" in heads
assert "The short version" in heads
assert "What we found" in heads
# the appendix heading is last among TOP-LEVEL sections; the questionnaire's own
# per-domain sub-headings follow it, which is correct nesting
appendix_idx = next(i for i, h in enumerate(heads) if "Appendix A" in h)
for later in heads[appendix_idx + 1:]:
    assert later not in ("For leadership", "For the technical team", "The short version",
                          "What we found"), f"{later} must not come after the appendix"
body = "\n".join(p.text for p in doc.paragraphs)
assert "Layout Eagle County layout v2" in body, "the report stamps which layout version produced it"
print("PASS: the Word export renders the SAME custom layout — custom titles, custom order, appendix last, "
      "and stamps the layout version")

# reordering is enough to move a section -- no code change
reordered = [
    {"type": "header", "title": "Reordered"},
    {"type": "questionnaire", "title": "Questionnaire FIRST"},
    {"type": "executive_summary"},
]
client.post("/api/v1/report-templates", json={"name": "Questionnaire-first", "blocks": reordered})
r = client.get(f"/api/v1/security-reviews/{rid}/export.docx")
doc = Document(io.BytesIO(r.content))
heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
assert heads[0] == "Reordered"
assert heads[1] == "Questionnaire FIRST" or any("Questionnaire FIRST" in h for h in heads[:3])
print("PASS: reordering the layout moves the section in the generated document — no code change, no deploy")

# reset restores the stock order as a new version
r = client.post("/api/v1/report-templates/reset")
assert r.json()["version"] == 4
rep = client.get(f"/api/v1/security-reviews/{rid}/report-data").json()
assert [b["type"] for b in rep["layout"]][-1] == "questionnaire"
assert rep["layout"][0]["title"] == "Security Review Report"
print("PASS: reset restores the stock layout as a new version, without destroying the custom ones")

# ============ hiding a section removes it everywhere ============

client.post("/api/v1/report-templates", json={"name": "No matrix", "blocks": [
    {"type": "header"}, {"type": "risk_verdict"},
    {"type": "risk_matrix", "visible": False}, {"type": "executive_summary"}]})
rep = client.get(f"/api/v1/security-reviews/{rid}/report-data").json()
assert "risk_matrix" not in [b["type"] for b in rep["layout"]]
r = client.get(f"/api/v1/security-reviews/{rid}/export.docx")
doc = Document(io.BytesIO(r.content))
heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
assert not any("Risk matrix" in h for h in heads)
print("PASS: hiding the risk matrix removes it from the report data AND the Word export")
