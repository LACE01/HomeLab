"""Security Reviews items 21-27 + 39: notes double-submit + rich text, matrix and
compensating controls in the report, Recommendation vs Decision, reviewer
reassignment, access-controlled sharing, and the .docx export."""
import os, sys, asyncio, uuid
os.environ["MONGO_URL"] = "mongodb://localhost:27017"
os.environ["DB_NAME"] = "test_sr_fixes"
os.environ["JWT_SECRET"] = "testsecret"
sys.path.insert(0, ".")

from mongomock_motor import AsyncMongoMockClient
import db as db_module
db_module.client = AsyncMongoMockClient()
db_module.db = db_module.client["test_sr_fixes"]

import server
import auth_utils
from routes import security_reviews as sr_route
sr_route.db = db_module.db

from fastapi.testclient import TestClient

admin_user = {"id": "u1", "email": "admin@x.com", "role": "admin", "name": "Admin", "teams": [], "team": "SecOps"}
app = server.app


def _as(u):
    """The shared-report endpoint takes get_current_user_optional, which calls
    get_current_user internally rather than through the dependency graph -- so
    the test harness has to override both to impersonate a user."""
    app.dependency_overrides[auth_utils.get_current_user] = lambda: u
    app.dependency_overrides[auth_utils.get_current_user_optional] = lambda: u


_as(admin_user)
client = TestClient(app)
db = db_module.db


def run(coro):
    return asyncio.get_event_loop().run_until_complete(coro)


run(db.users.insert_many([
    {"id": "u1", "email": "admin@x.com", "name": "Admin", "role": "admin"},
    {"id": "u2", "email": "analyst@x.com", "name": "Ana Lyst", "role": "admin"},
]))

client.get("/api/v1/security-reviews/meta")
review = client.post("/api/v1/security-reviews", json={
    "title": "Acme SaaS", "review_type": "New software purchase (SaaS / COTS / on-prem)",
    "entity_name": "Acme", "entity_domain": "acme.com",
    "data_classifications": ["PII (Colorado)"],
}).json()
rid = review["id"]

# ============ 21: notes double-submit ============

r = client.post(f"/api/v1/security-reviews/{rid}/notes", json={"text": "Called the vendor"})
assert r.status_code == 200
first_id = r.json()["id"]
r2 = client.post(f"/api/v1/security-reviews/{rid}/notes", json={"text": "Called the vendor"})
assert r2.json()["id"] == first_id, "identical note within 5s must return the existing row, not insert a duplicate"
assert run(db.security_review_notes.count_documents({"review_id": rid})) == 1
# a genuinely different note still goes through
r3 = client.post(f"/api/v1/security-reviews/{rid}/notes", json={"text": "Second, different note"})
assert r3.json()["id"] != first_id
assert run(db.security_review_notes.count_documents({"review_id": rid})) == 2
print("PASS: identical rapid note submissions collapse to one row (double-submit fix), distinct notes still save")

# ============ 22: rich text ============

r = client.post(f"/api/v1/security-reviews/{rid}/notes", json={
    "text": "Bold and code", "html": "<b>Bold</b> and <pre>code</pre>"})
assert r.json()["html"] == "<b>Bold</b> and <pre>code</pre>"
items = client.get(f"/api/v1/security-reviews/{rid}/notes").json()["items"]
assert any(i.get("html") for i in items)
print("PASS: notes persist rich-text HTML alongside the plain-text copy")

# ============ 25: reviewer reassignment ============

r = client.get("/api/v1/security-reviews/assignable-users")
emails = {u["email"] for u in r.json()["items"]}
assert "analyst@x.com" in emails
r = client.post(f"/api/v1/security-reviews/{rid}/reassign", json={"assignee": "analyst@x.com"})
assert r.status_code == 200
doc = run(db.security_reviews.find_one({"id": rid}, {"_id": 0}))
assert doc["assignee"] == "analyst@x.com"
audit = client.get(f"/api/v1/security-reviews/{rid}/audit").json()["items"]
assert any(a["action"] == "reviewer_reassigned" for a in audit)
r = client.post(f"/api/v1/security-reviews/{rid}/reassign", json={"assignee": "ghost@x.com"})
assert r.status_code == 404
print("PASS: reviewer reassignment validates the target user, updates the review, and writes an audit entry")

# ============ 24: recommendation, distinct from decision ============

r = client.put(f"/api/v1/security-reviews/{rid}/recommendation", json={
    "what_was_reviewed": "Acme scheduling SaaS", "why": "Replace paper process",
    "recommendation": "Approve with conditions", "rationale": "Residual is acceptable with MFA enforced"})
assert r.status_code == 200
client.put(f"/api/v1/security-reviews/{rid}/decision", json={
    "outcome": "Rejected", "decision_maker": "CISO", "rationale": "Budget"})
doc = run(db.security_reviews.find_one({"id": rid}, {"_id": 0}))
assert doc["recommendation"]["recommendation"] == "Approve with conditions"
assert doc["decision"]["outcome"] == "Rejected"
assert doc["recommendation"]["authored_by"] == "admin@x.com"
print("PASS: Recommendation and Decision are stored separately -- a decision that diverges doesn't overwrite the recommendation")

# ============ 23 + 27: matrix points + compensating controls in report data ============

client.put(f"/api/v1/security-reviews/{rid}/risk-score", json={
    "inherent_likelihood": 4,
    "inherent_impacts": {"confidentiality": 5, "integrity": 3, "availability": 2, "compliance_legal": 4, "reputational": 3},
    "compensating_controls": "Enforce SSO+MFA via IdP; restrict data residency to US",
    "residual_likelihood": 2,
    "residual_impacts": {"confidentiality": 3, "integrity": 2, "availability": 2, "compliance_legal": 2, "reputational": 2},
    "not_adopting_likelihood": 3, "not_adopting_impacts": {"confidentiality": 1, "integrity": 1, "availability": 4, "compliance_legal": 1, "reputational": 2},
})
rep = client.get(f"/api/v1/security-reviews/{rid}/report-data").json()
assert rep["compensating_controls"].startswith("Enforce SSO+MFA")
pts = {p["label"]: p for p in rep["matrix_points"]}
assert pts["Inherent"] == {"label": "Inherent", "likelihood": 4, "impact": 5, "band": "Critical"}
assert pts["Residual"]["band"] == "Medium"
assert pts["Not adopting"]["impact"] == 4
assert rep["recommendation"]["recommendation"] == "Approve with conditions"
print("PASS: report data carries the 5x5 matrix points (inherent/residual/not-adopting), compensating controls, and the recommendation")

# ============ 26: access-controlled sharing ============

r = client.post(f"/api/v1/security-reviews/{rid}/share", json={})
assert r.status_code == 400, "must specify exactly one share mode"
r = client.post(f"/api/v1/security-reviews/{rid}/share",
                 json={"email": "ext@partner.com", "platform_user_email": "analyst@x.com"})
assert r.status_code == 400

# --- external email + one-time code ---
r = client.post(f"/api/v1/security-reviews/{rid}/share", json={"email": "ext@partner.com"})
assert r.status_code == 200
token = r.json()["token"]
assert r.json()["mode"] == "email_code"
grant = run(db.security_review_share_grants.find_one({"token": token}, {"_id": 0}))
code = grant["code"]
assert code and len(code) == 6

m = client.get(f"/api/v1/shared/security-review/{token}/meta").json()
assert m["requires_code"] is True
assert m["recipient_hint"] == "ex*@partner.com"
assert "review" not in m, "meta must not leak report content before verification"

# the link ALONE must not work
r = client.get(f"/api/v1/shared/security-review/{token}")
assert r.status_code == 401, "anyone-with-the-link access must be gone"

r = client.post(f"/api/v1/shared/security-review/{token}/verify", json={"code": "000000"})
assert r.status_code == 403
r = client.post(f"/api/v1/shared/security-review/{token}/verify", json={"code": code})
assert r.status_code == 200
payload = r.json()
assert payload["review"]["review_number"] == review["review_number"]
assert payload["matrix_points"] and payload["compensating_controls"]
# after verification the link works for that recipient
r = client.get(f"/api/v1/shared/security-review/{token}")
assert r.status_code == 200
grant = run(db.security_review_share_grants.find_one({"token": token}, {"_id": 0}))
assert grant["verified"] is True and grant["view_count"] >= 2
print("PASS: emailed-code sharing gates the report -- the link alone 401s, wrong codes 403, the right code unlocks it and views are counted")

# lockout after 5 bad codes
r2 = client.post(f"/api/v1/security-reviews/{rid}/share", json={"email": "other@partner.com"})
t2 = r2.json()["token"]
for _ in range(5):
    client.post(f"/api/v1/shared/security-review/{t2}/verify", json={"code": "111111"})
r = client.post(f"/api/v1/shared/security-review/{t2}/verify", json={"code": "111111"})
assert r.status_code == 429
print("PASS: repeated wrong codes lock the grant (brute-force guard)")

# --- platform user mode ---
r = client.post(f"/api/v1/security-reviews/{rid}/share", json={"platform_user_email": "analyst@x.com"})
assert r.status_code == 200
ptoken = r.json()["token"]
r = client.post(f"/api/v1/security-reviews/{rid}/share", json={"platform_user_email": "nobody@x.com"})
assert r.status_code == 404, "can't share with a non-existent platform user"

# admin (not the grantee) is refused; the grantee gets in
r = client.get(f"/api/v1/shared/security-review/{ptoken}")
assert r.status_code == 403
_as({"id": "u2", "email": "analyst@x.com", "role": "admin", "name": "Ana", "teams": []})
r = client.get(f"/api/v1/shared/security-review/{ptoken}")
assert r.status_code == 200
_as(admin_user)
print("PASS: platform-user sharing admits only the granted user, not merely anyone signed in")

# --- revocation ---
grants = client.get(f"/api/v1/security-reviews/{rid}/shares").json()["items"]
assert all("code" not in g and "token" not in g for g in grants), "codes/tokens must never be listed back out"
gid = next(g["id"] for g in grants if g["recipient"] == "ext@partner.com")
client.delete(f"/api/v1/security-reviews/{rid}/shares/{gid}")
r = client.get(f"/api/v1/shared/security-review/{token}")
assert r.status_code == 404
print("PASS: share grants are listable without leaking secrets and revocation kills access immediately")

# ============ 39: docx export ============

r = client.get(f"/api/v1/security-reviews/{rid}/export.docx")
assert r.status_code == 200
assert r.headers["content-type"].startswith("application/vnd.openxmlformats")
assert review["review_number"] in r.headers["content-disposition"]
blob = r.content
assert len(blob) > 5000
import io
from docx import Document
doc = Document(io.BytesIO(blob))
heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
for expected in ("Security Review Report", "Risk verdict", "Executive summary",
                  "Reviewer recommendation", "Decision"):
    assert any(expected in h for h in heads), f"missing section {expected}: {heads}"
assert any("Compensating controls" in h for h in heads)
assert any("Risk matrix" in h for h in heads)
# the matrix table is a real 6x6 Word table, not an image
matrix_tables = [t for t in doc.tables if len(t.rows) == 6 and len(t.columns) == 6]
assert matrix_tables, "5x5 matrix must render as an editable Word table"
all_text = "\n".join(p.text for p in doc.paragraphs)
assert "Enforce SSO+MFA" in all_text
assert "Approve with conditions" in all_text
assert "differs from the reviewer's recommendation" in all_text  # decision was Rejected
print("PASS: .docx export produces a style-mapped editable Word document with every report section, "
      "the matrix as a real table, compensating controls, and a recommendation-vs-decision divergence note")
