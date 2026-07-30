"""Follow-up batch: External Checks panel-merge fix + executive framing, review
asset linking (individual / by team / by tag), supporting-document attachments,
a complete two-part report, and OpenCTI as a Threat News source."""
import os, sys, asyncio, uuid, io
os.environ["MONGO_URL"] = "mongodb://localhost:27017"
os.environ["DB_NAME"] = "test_sr_completeness"
os.environ["JWT_SECRET"] = "testsecret"
sys.path.insert(0, ".")

from mongomock_motor import AsyncMongoMockClient
import db as db_module
db_module.client = AsyncMongoMockClient()
db_module.db = db_module.client["test_sr_completeness"]

import server
import auth_utils
from routes import security_reviews as sr_route
from routes import cti as cti_route
sr_route.db = db_module.db
cti_route.db = db_module.db
import external_checks
import cti

from fastapi.testclient import TestClient

admin_user = {"id": "u1", "email": "admin@x.com", "role": "admin", "name": "Admin", "teams": []}
app = server.app
app.dependency_overrides[auth_utils.get_current_user] = lambda: admin_user
app.dependency_overrides[auth_utils.get_current_user_optional] = lambda: admin_user
client = TestClient(app)
db = db_module.db


def run(coro):
    return asyncio.get_event_loop().run_until_complete(coro)


class Resp:
    """Faithful httpx-like response: .text is real JSON and content-type is set,
    so anything inspecting the response shape sees a well-formed reply."""

    def __init__(self, status_code=200, json_data=None, text=None, headers=None, url="https://acme.com/"):
        import json as _json
        self.status_code = status_code
        self._json = json_data
        self.text = text if text is not None else (
            _json.dumps(json_data) if json_data is not None else "")
        default_headers = {"Content-Type": "application/json"} if json_data is not None else {}
        self.headers = headers if headers is not None else default_headers
        self.url = url
    def json(self):
        if self._json is None:
            raise ValueError("no json")
        return self._json


class FakeHttp:
    def __init__(self, *a, **kw): pass
    async def __aenter__(self): return self
    async def __aexit__(self, *a): return False
    async def get(self, url, **kw):
        if "opencorporates" in url:
            return Resp(json_data={"results": {"companies": [{"company": {
                "name": "ACME CORPORATION", "jurisdiction_code": "us_co",
                "current_status": "Good Standing", "incorporation_date": "2009-04-01"}}]}})
        if "nvd.nist.gov" in url:
            return Resp(json_data={"totalResults": 0, "vulnerabilities": []})
        if "crt.sh" in url or "certspotter" in url:
            return Resp(json_data=[])
        if "shodan.io" in url:
            return Resp(json_data={"total": 0, "matches": []})
        return Resp(headers={"Strict-Transport-Security": "max-age=1"})
    async def post(self, url, **kw):
        q = (kw.get("json") or {}).get("query", "")
        if "reports(" in q:
            return Resp(json_data={"data": {"reports": {"edges": [
                {"node": {"id": "rep-1", "name": "Acme Corp breach analysis",
                          "description": "Detailed writeup of the Acme Corp incident.",
                          "published": "2026-07-20T00:00:00Z",
                          "createdBy": {"name": "CTI Team"},
                          "objectLabel": [{"value": "ransomware"}],
                          "externalReferences": {"edges": [
                              {"node": {"url": "https://cti.example/report/1", "source_name": "internal"}}]}}},
                {"node": {"id": "rep-2", "name": "Unrelated malware family roundup",
                          "description": "Nothing to do with us.",
                          "published": "2026-07-21T00:00:00Z",
                          "createdBy": None, "objectLabel": [],
                          "externalReferences": {"edges": []}}},
            ]}}})
        return Resp(json_data={"data": {}})


import httpx
_real = httpx.AsyncClient

client.get("/api/v1/security-reviews/meta")
review = client.post("/api/v1/security-reviews", json={
    "title": "Acme SaaS", "review_type": "New software purchase (SaaS / COTS / on-prem)",
    "entity_name": "Acme", "entity_domain": "acme.com",
    "data_classifications": ["PII (Colorado)"]}).json()
rid = review["id"]


# =========================================================================
# External Checks: running one panel must not wipe the other
# =========================================================================

httpx.AsyncClient = FakeHttp
r = client.post(f"/api/v1/security-reviews/{rid}/external-checks", params={"panel": "company"})
assert r.status_code == 200
assert "company_posture" in r.json() and "technical_posture" not in r.json()

r = client.post(f"/api/v1/security-reviews/{rid}/external-checks", params={"panel": "technical"})
payload = r.json()
# THE BUG: this used to replace the stored document, wiping the company panel.
assert "technical_posture" in payload
assert "company_posture" in payload, "running one panel must not wipe the other"
assert sorted(payload["panels_present"]) == ["company_posture", "technical_posture"]

stored = run(db.security_reviews.find_one({"id": rid}, {"_id": 0}))["external_checks"]
assert stored.get("company_posture") and stored.get("technical_posture")
print("PASS: running a single External Checks panel merges into the stored result instead of wiping the other panel")

# re-running one panel refreshes only that panel
before_company_ran = stored["company_posture"]["ran_at"]
r = client.post(f"/api/v1/security-reviews/{rid}/external-checks", params={"panel": "technical"})
after = run(db.security_reviews.find_one({"id": rid}, {"_id": 0}))["external_checks"]
assert after["company_posture"]["ran_at"] == before_company_ran, "the untouched panel must not be re-run"
print("PASS: re-running one panel leaves the other panel's results and timestamp untouched")


# =========================================================================
# Executive framing
# =========================================================================

company = {c["check"]: c for c in after["company_posture"]["results"]}
reg = company["corporate_registration"]
assert reg["label"] == "Is this a real, registered company?"
assert reg["what_it_means"] and reg["why_it_matters"]
assert reg["status_plain"] in ("No concerns found", "Worth a look", "Couldn't check automatically", "Not set up yet")
assert all(c.get("label") and c.get("why_it_matters") for c in after["technical_posture"]["results"])
print("PASS: every check carries a plain-English question, what-we-checked, why-it-matters, and a readable status")

summary = after["company_posture"]["summary"]
assert summary["verdict"] in ("ok", "attention", "partial")
assert len(summary["headline"]) > 40 and "checks on the vendor as a company" in summary["headline"]
assert set(summary["counts"]) == {"attention", "ok", "unchecked"}
print("PASS: each panel produces a one-sentence executive verdict, so a director doesn't have to read eleven cards")

httpx.AsyncClient = _real


# =========================================================================
# OpenCorporates is configurable
# =========================================================================

import seed
names = {c["name"] for c in seed.INTEGRATION_CATALOG} if hasattr(seed, "INTEGRATION_CATALOG") else set()
if not names:
    src = open("seed.py").read()
    assert '"name": "OpenCorporates"' in src
else:
    assert "OpenCorporates" in names
print("PASS: OpenCorporates now appears in the Integrations catalog so its API token can be configured")


# =========================================================================
# Asset linking: individual, by team, by tag
# =========================================================================

run(db.assets.insert_many([
    {"id": "a1", "hostname": "web01", "ip": "10.0.0.1", "owner_team": "IT", "criticality": "High",
     "tags": ["public", "prod"]},
    {"id": "a2", "hostname": "web02", "ip": "10.0.0.2", "owner_team": "IT", "criticality": "Medium",
     "tags": ["prod"]},
    {"id": "a3", "hostname": "hr01", "ip": "10.0.0.3", "owner_team": "HR", "criticality": "Low",
     "tags": ["internal"]},
    {"id": "a4", "hostname": "kiosk1", "ip": "10.0.0.4", "owner_team": None, "criticality": "Low",
     "tags": ["public"]},
]))
run(db.findings.insert_many([
    {"id": "f1", "asset_id": "a1", "status": "New", "severity": "Critical", "title": "x"},
    {"id": "f2", "asset_id": "a1", "status": "Valid", "severity": "Low", "title": "y"},
    {"id": "f3", "asset_id": "a2", "status": "Fixed validated", "severity": "High", "title": "z"},
]))

r = client.get("/api/v1/security-reviews/asset-picker")
picker = r.json()
assert picker["teams"] == ["HR", "IT"]
assert set(picker["tags"]) == {"public", "prod", "internal"}
r = client.get("/api/v1/security-reviews/asset-picker", params={"q": "web"})
assert {a["hostname"] for a in r.json()["items"]} == {"web01", "web02"}
print("PASS: the asset picker searches by hostname/IP and lists the distinct teams and tags available for bulk selection")

r = client.post(f"/api/v1/security-reviews/{rid}/assets", json={"asset_ids": ["a3"]})
assert r.json()["added"] == 1 and r.json()["linked_total"] == 1

r = client.post(f"/api/v1/security-reviews/{rid}/assets", json={"teams": ["IT"]})
assert r.json()["matched_by"]["teams"]["IT"] == 2
assert r.json()["linked_total"] == 3, "team add must APPEND to the individually-picked asset"

r = client.post(f"/api/v1/security-reviews/{rid}/assets", json={"tags": ["public"]})
# a1 already linked via the team; only a4 is new
assert r.json()["added"] == 1 and r.json()["linked_total"] == 4
print("PASS: assets link individually, in bulk by team, and in bulk by tag — combining without duplicating")

r = client.post(f"/api/v1/security-reviews/{rid}/assets", json={"asset_ids": ["a1"], "replace": True})
assert r.json()["linked_total"] == 1
r = client.post(f"/api/v1/security-reviews/{rid}/assets", json={"teams": ["Nonexistent"]})
assert r.status_code == 400, "a selector matching nothing should say so rather than silently no-op"
print("PASS: replace mode resets the scope, and a selector that matches nothing is rejected")

client.post(f"/api/v1/security-reviews/{rid}/assets", json={"teams": ["IT"], "tags": ["internal"]})
r = client.get(f"/api/v1/security-reviews/{rid}/assets")
linked = r.json()
assert linked["total"] == 3
web01 = next(a for a in linked["items"] if a["hostname"] == "web01")
assert web01["open_findings"] == 2 and web01["critical_high_findings"] == 1
assert linked["items"][0]["hostname"] == "web01", "worst-first ordering"
print("PASS: linked assets report their open and Critical/High finding counts, worst first")

r = client.post(f"/api/v1/security-reviews/{rid}/assets/unlink", json={"asset_ids": ["a3"]})
assert r.json()["linked_total"] == 2
audit = client.get(f"/api/v1/security-reviews/{rid}/audit").json()["items"]
assert any(a["action"] == "assets_linked" for a in audit)
assert any(a["action"] == "assets_unlinked" for a in audit)
print("PASS: unlinking works and both link and unlink are audited")

# linked assets feed the existing environment-health hook
r = client.get(f"/api/v1/security-reviews/{rid}/autofill/open_findings_pull")
assert r.json()["total_open"] == 2 and r.json()["asset_count"] == 2
print("PASS: linked assets feed the open_findings_pull auto-fill hook (scope drives environment health)")


# =========================================================================
# Supporting documents
# =========================================================================

tiny = "data:application/pdf;base64," + "A" * 100
r = client.post(f"/api/v1/security-reviews/{rid}/attachments", json={
    "name": "SOC2-2026.pdf", "mime": "application/pdf", "data_url": tiny,
    "category": "certificate", "description": "Current SOC 2 Type II"})
assert r.status_code == 200
att = r.json()
assert "data_url" not in att, "the create response must not echo the whole file back"
assert att["category"] == "certificate" and att["size_bytes"] > 0

r = client.post(f"/api/v1/security-reviews/{rid}/attachments", json={
    "name": "bad.pdf", "data_url": "not-a-data-url"})
assert r.status_code == 400
r = client.post(f"/api/v1/security-reviews/{rid}/attachments", json={
    "name": "huge.pdf", "data_url": "data:application/pdf;base64," + "A" * 14_000_001})
assert r.status_code == 413
print("PASS: supporting documents attach to the review, validate the payload, and enforce a size limit")

r = client.get(f"/api/v1/security-reviews/{rid}/attachments")
assert len(r.json()["items"]) == 1
client.delete(f"/api/v1/security-reviews/{rid}/attachments/{att['id']}")
assert len(client.get(f"/api/v1/security-reviews/{rid}/attachments").json()["items"]) == 0
# re-add for the report test
client.post(f"/api/v1/security-reviews/{rid}/attachments", json={
    "name": "SOC2-2026.pdf", "data_url": tiny, "category": "certificate"})
print("PASS: attachments list and delete cleanly")


# =========================================================================
# Editable findings
# =========================================================================

f = client.post(f"/api/v1/security-reviews/{rid}/findings", json={
    "description": "No MFA", "severity": "High"}).json()
r = client.patch(f"/api/v1/security-reviews/{rid}/findings/{f['id']}", json={
    "description": "No MFA available below Enterprise tier",
    "severity": "Critical", "recommendation": "Buy Enterprise; enforce via IdP",
    "owner": "IT Ops", "due_date": "2026-10-01", "affected_component": "Admin console",
    "cis_mapping": "6.5", "is_condition_of_approval": True, "condition_deadline": "2026-09-15"})
assert r.status_code == 200
upd = r.json()
assert upd["severity"] == "Critical" and upd["owner"] == "IT Ops"
assert upd["affected_component"] == "Admin console" and upd["cis_mapping"] == "6.5"
assert upd["is_condition_of_approval"] is True and upd["condition_deadline"] == "2026-09-15"
print("PASS: findings are fully editable after creation (description, severity, recommendation, owner, dates, condition)")


# =========================================================================
# Complete two-part report
# =========================================================================

client.put(f"/api/v1/security-reviews/{rid}/risk-score", json={
    "inherent_likelihood": 4,
    "inherent_impacts": {"confidentiality": 5, "integrity": 3, "availability": 2,
                          "compliance_legal": 4, "reputational": 3},
    "compensating_controls": "Enforce SSO+MFA; restrict data residency to US",
    "residual_likelihood": 2,
    "residual_impacts": {"confidentiality": 3, "integrity": 2, "availability": 2,
                          "compliance_legal": 2, "reputational": 2}})
client.put(f"/api/v1/security-reviews/{rid}/recommendation", json={
    "what_was_reviewed": "Acme scheduling SaaS", "why": "Replace paper process",
    "recommendation": "Approve with conditions", "rationale": "Acceptable post-controls"})
client.post(f"/api/v1/security-reviews/{rid}/notes", json={
    "text": "Called the vendor about MFA", "html": "<b>Called</b> the vendor about MFA"})
client.put(f"/api/v1/security-reviews/{rid}/decision", json={
    "outcome": "Approved with Conditions", "decision_maker": "CISO"})
# answer one questionnaire item so the technical appendix has content to render
_applicable = client.get(f"/api/v1/security-reviews/{rid}").json()["applicable_questions"]
client.put(f"/api/v1/security-reviews/{rid}/responses", json={
    "question_order": _applicable[0]["order"], "answer": "no",
    "evidence_text": "Vendor confirmed this isn't available"})

rep = client.get(f"/api/v1/security-reviews/{rid}/report-data").json()
assert rep["notes"] and rep["notes"][0]["html"]
assert rep["attachments"] and rep["attachments"][0]["name"] == "SOC2-2026.pdf"
assert "data_url" not in rep["attachments"][0], "report data must not carry file bytes"
assert rep["linked_assets"] and len(rep["linked_assets"]) == 2
assert rep["external_checks"]["company_posture"] and rep["external_checks"]["technical_posture"]
assert rep["compensating_controls"] and rep["matrix_points"]
assert rep["recommendation"]["recommendation"] == "Approve with conditions"
assert rep["questionnaire_scoring"] is not None and "confidence_pct" in rep["questionnaire_scoring"]
print("PASS: report data now carries EVERYTHING — notes, attachments, in-scope assets, external checks, "
      "recommendation, controls, matrix, and the confidence read")

# the shared/external copy deliberately excludes internal notes
r = client.post(f"/api/v1/security-reviews/{rid}/share", json={"email": "ext@partner.com"})
token = r.json()["token"]
grant = run(db.security_review_share_grants.find_one({"token": token}, {"_id": 0}))
shared = client.post(f"/api/v1/shared/security-review/{token}/verify",
                      json={"code": grant["code"]}).json()
assert shared["notes"] == [], "internal working notes must never leave in a shared report"
assert shared["linked_assets"] and shared["external_checks"], "but scope and verification results should"
print("PASS: the shared copy includes scope and verification results but never the internal working notes")

# docx carries both parts
r = client.get(f"/api/v1/security-reviews/{rid}/export.docx")
assert r.status_code == 200
from docx import Document
doc = Document(io.BytesIO(r.content))
heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
assert any("Part 1 — Executive summary" in h for h in heads)
assert any("Part 2 — Technical detail" in h for h in heads)
# section titles now come from the report LAYOUT template, so match on the
# stock titles it ships with
for expected in ("In-scope assets", "External verification checks", "Analyst working notes",
                  "Supporting documents", "questionnaire responses"):
    assert any(expected.lower() in h.lower() for h in heads), f"missing {expected}: {heads}"
body = "\n".join(p.text for p in doc.paragraphs)
# assets render as a real Word TABLE (editable), so check table cells too
table_text = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
assert "web01" in table_text, "in-scope assets must render as an editable table"
assert "SOC2-2026.pdf" in body
assert "Called the vendor about MFA" in body, "rich-text notes must flatten to readable text in Word"
assert "<b>" not in body, "HTML markup must not leak into the Word document"
print("PASS: the Word export is split into an executive part and a technical part and contains assets, "
      "checks, notes (HTML flattened), documents, and the questionnaire")


# =========================================================================
# OpenCTI as a Threat News source
# =========================================================================

run(db.vendors.insert_one({"id": "v1", "name": "Acme Corp", "domain": "acme.com"}))
run(db.integrations.insert_one({"id": "octi", "name": "OpenCTI", "type": "threat_intel",
                                 "config": {"endpoint": "https://opencti.example", "api_key": "tok"}}))
httpx.AsyncClient = FakeHttp

run(db.cti_articles.delete_many({}))
r = client.post("/api/v1/cti/opencti/sync")
assert r.status_code == 200, r.text
res = r.json()
assert res["reports_seen"] == 2 and res["articles_created"] == 2
assert res["articles_matched"] == 1, "only the Acme report should match a tracked vendor"
arts = run(db.cti_articles.find({"feed_id": "opencti"}, {"_id": 0}).to_list(10))
acme = next(a for a in arts if "Acme" in a["title"])
assert acme["source"].startswith("OpenCTI") and "CTI Team" in acme["source"]
assert acme["opencti_report_id"] == "rep-1"
assert acme["labels"] == ["ransomware"]
assert acme["matches"][0]["kind"] == "vendor"
assert run(db.security_events.count_documents({"event_type": "threat_news_match"})) >= 1

# dedupes on re-sync
r = client.post("/api/v1/cti/opencti/sync")
assert r.json()["articles_created"] == 0
httpx.AsyncClient = _real
print("PASS: OpenCTI Reports flow into the same Threat News stream — matched against vendors/domains/keywords, "
      "raising the same alerts, and deduped on re-sync")

# and they show up in the normal article list
r = client.get("/api/v1/cti/articles", params={"matched_only": True})
assert any(a.get("feed_id") == "opencti" for a in r.json()["items"])
print("PASS: OpenCTI articles appear alongside RSS articles in the Threat News list")
