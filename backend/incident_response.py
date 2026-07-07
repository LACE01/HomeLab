"""Incident Response module -- a helpdesk-facing triage wizard that scores a
reported problem across common attack-vector categories, an admin-configurable IR
plan (phases/roles/classification/tool catalog), and a single collaborative IR case
("ticket") that consolidates the timeline, evidence, and closure report for one
incident from open to close.

This was modeled on a real municipal IR plan the app's operator shared (a 14-step
phased response process, a severity classification matrix, an Incident Handling
Ledger, and a set of closure report templates) but everything here is written from
scratch and kept generic/paraphrased on purpose -- the whole point of this module is
that *any* org can reshape the phases, roles, classification thresholds, wizard
questions, and tool catalog to match how *their* org actually runs incident response,
not just replay one county's plan verbatim. All of the DEFAULT_* structures below are
just a sensible starting point (seeded once, on first use) -- every one of them lives
in its own admin-editable config document, same philosophy as rbac.py/criticality.py.

Architecture, briefly:
  - IR plan config (phases, roles, classification levels) -- singleton docs, editable
    under Admin -> Incident Response Setup. Phases can be renamed/reordered/added/
    removed; nothing about the 14-step structure is hardcoded into behavior, it's
    just the seeded starting content.
  - Wizard config (categories, questions, per-category action plans) -- also a
    singleton doc, admin-editable. Scoring is a simple additive weights model (not a
    branching decision tree): every question is shown once, each answer option
    contributes weight to one or more outcome categories (and optionally a
    "severity_points" contribution that drives classification, and an
    "immediate_containment" flag for the ransomware/active-attack fast path). This
    keeps the config a flat list an admin can actually edit in a table, instead of a
    branching tree that's hard to reason about or keep in sync.
  - Tool/resource catalog -- a real collection (db.ir_tools), CRUD, each entry tagged
    with which outcome categories it's relevant to, surfaced in the wizard's result
    screen and the resulting case.
  - IR case -- the single ticket, with a phase checklist (instantiated from the plan
    config at case-open time so later plan edits don't retroactively rewrite an
    in-progress case's checklist), a unified timeline (db.ir_case_events -- notes,
    screenshots, status/phase changes, evidence, role assignments all in one feed),
    an evidence manifest, and a closure report.
  - Collaboration: deliberately no new realtime infra -- the case page polls/refreshes
    like the rest of the app, and the timeline itself *is* the "who did what when"
    record multiple simultaneous responders need. Optimistic, not locking.
  - Google Sheets: one-way, best-effort export. Rather than requiring full Google
    OAuth/service-account plumbing, this posts each new timeline event to a Google
    Apps Script Web App URL the org deploys themselves (a few lines of Apps Script
    bound to their sheet that appends a row per POST) -- zero new Python
    dependencies, no OAuth flow, and it's a well-established lightweight pattern for
    "push data into Sheets without the full API". Configured under Integrations ->
    Google Sheets (default target for all cases) with an optional per-case override
    for orgs that want one sheet per incident.
"""
import uuid
from datetime import datetime, timezone


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _id() -> str:
    return str(uuid.uuid4())


# --------------------------------------------------------------------------------
# Classification levels (paraphrased generic version of a severity/response matrix --
# any org's actual thresholds will differ, this is just a reasonable starting point).
# --------------------------------------------------------------------------------
CLASSIFICATION_LEVELS = ["Critical", "Significant", "Moderate", "Minor", "Negligible"]

# Maps an IR classification to the severity vocabulary the rest of the app's
# notification rules already filter on (finding severities) -- reusing that instead
# of inventing a second severity taxonomy just for IR alerts.
CLASSIFICATION_TO_SEVERITY = {
    "Critical": "Critical", "Significant": "High", "Moderate": "Medium",
    "Minor": "Low", "Negligible": "Info",
}


def _default_classification_config() -> dict:
    return {
        "levels": [
            {"level": "Critical", "criteria": "Targeted attack causing a major outage, or confirmed loss/exposure of sensitive data.",
             "response": "Assign and alert the Incident Commander, on-call team, and management immediately. Engage a third-party IR vendor if one is available. Proceed through the full incident response process."},
            {"level": "Significant", "criteria": "Unauthorized changes to a system (e.g. defacement), or a suspected-but-unconfirmed exposure of sensitive data.",
             "response": "Assign and alert the Incident Commander and on-call team. Proceed through the full incident response process."},
            {"level": "Moderate", "criteria": "Suspected compromise where the impact isn't yet known, multiple reports of unusual/unexpected behavior, or a malware outbreak.",
             "response": "Assign and alert the Incident Commander and on-call team. Proceed through the full incident response process."},
            {"level": "Minor", "criteria": "An unsuccessful attack attempt, or an isolated event with contained, well-understood impact.",
             "response": "Monitor for further impact and handle through normal operational process. Formal incident response isn't required unless the situation changes."},
            {"level": "Negligible", "criteria": "An isolated alert, spam email, or false positive.",
             "response": "No formal response needed -- log and close."},
        ],
        "updated_at": _now_iso(),
    }


# --------------------------------------------------------------------------------
# Default IR plan phases -- a generic paraphrase of a common phased IR process
# (initial notification -> triage/classify -> command & briefing -> assign roles ->
# define objectives -> plan -> contain -> preserve evidence -> analyze -> communicate
# -> restore -> root cause -> final report -> post-mortem), NOT tied to any specific
# org's exact wording. Every phase is fully editable; admins can add/remove/reorder.
# --------------------------------------------------------------------------------
def _default_phases() -> list:
    phases = [
        {"name": "Notify & Begin Intake", "responsible_party": "First Responder / Helpdesk",
         "tasks": ["Escalate the suspected incident to the Triage Lead immediately",
                    "Open the incident ledger and record the current time and known event timestamps",
                    "Gather what's known: how it was found, what systems/data may be involved, who else is aware",
                    "Collect and preserve any evidence already gathered by the reporter"],
         "objectives": ["Triage Lead is aware of the incident", "All known background/evidence is captured in the case"],
         "things_needed": ["Incident response kit / evidence bags", "Time-stamp reference"]},
        {"name": "Immediate Containment (active malware/ransomware only)", "responsible_party": "IT Operations / Security",
         "tasks": ["Identify which systems are impacted and isolate them from the network immediately",
                    "Do not power off affected systems where memory evidence may still be needed"],
         "objectives": ["Active spread is contained"],
         "things_needed": [], "parallel_fast_path": True},
        {"name": "Verify & Classify", "responsible_party": "Triage Lead",
         "tasks": ["Review the intake details (indicators, systems/parties involved, timeline, disclosures)",
                    "Classify the incident using the Classification levels (Admin -> Incident Response Setup)",
                    "Engage a third-party IR vendor if the classification calls for it"],
         "objectives": ["Incident is confirmed and classified", "Incident Commander is assigned and alerted"],
         "things_needed": []},
        {"name": "Incident Commander Takes Charge & Briefs", "responsible_party": "Incident Commander",
         "tasks": ["Transfer ownership from the Triage Lead to the Incident Commander",
                    "Schedule an initial briefing within minutes of classification",
                    "Establish where incident data/communication will live and record it in the case",
                    "Brief management using the current state, classification, and immediate needs"],
         "objectives": ["Management is briefed", "Case ownership is transferred to the IC", "On-call resources are engaged"],
         "things_needed": ["Conferencing/communication channel", "Org chart / phone tree"]},
        {"name": "Assign & Communicate Roles", "responsible_party": "Incident Commander",
         "tasks": ["Assign the mandatory roles (IT Ops/Security rep, Cybersecurity rep, Legal/Compliance rep)",
                    "Assign optional roles as needed (PIO, Safety Officer, Liaison, Section Chiefs)",
                    "Communicate assignments to everyone involved"],
         "objectives": ["Every needed role has an assigned, informed owner"], "things_needed": []},
        {"name": "Define Objectives", "responsible_party": "Incident Commander",
         "tasks": ["Define specific, near-term objectives for the current operational period",
                    "Record objectives in the Incident Action Plan"],
         "objectives": ["Objectives for this operational period are documented"], "things_needed": []},
        {"name": "Strategy & Tactics: Build the Incident Action Plan", "responsible_party": "Incident Commander + IR staff",
         "tasks": ["For each objective, define the strategy (approach) and specific tactics (actions)",
                    "Assign an owner and due date to each tactic",
                    "Track status (not started / in process / complete) as work proceeds"],
         "objectives": ["An actionable Incident Action Plan exists and is being worked"], "things_needed": []},
        {"name": "Contain & Isolate", "responsible_party": "IT Operations / Security",
         "tasks": ["Execute containment actions from the Incident Action Plan",
                    "Isolate affected systems, accounts, or network segments as planned"],
         "objectives": ["Further damage/spread is stopped"], "things_needed": []},
        {"name": "Preserve Evidence", "responsible_party": "IT Operations / Security",
         "tasks": ["Collect logs, images, and other evidence before it's lost or overwritten",
                    "Log every item in the evidence manifest with chain-of-custody details",
                    "Store evidence in the agreed-upon location"],
         "objectives": ["Evidence needed for analysis (and any legal process) is preserved"], "things_needed": []},
        {"name": "Analyze Evidence", "responsible_party": "Cybersecurity rep / IR staff",
         "tasks": ["Analyze preserved evidence to determine root cause, scope, and timeline",
                    "Separate causes from symptoms/effects (ask \"why\" repeatedly until you reach the real cause)"],
         "objectives": ["Root cause and full scope are understood"], "things_needed": []},
        {"name": "Create Communication Plan", "responsible_party": "Incident Commander / PIO",
         "tasks": ["Decide who needs to be told what, and when (internal staff, leadership, affected users, regulators, public)",
                    "Prepare any required breach/regulatory notifications"],
         "objectives": ["All required parties are notified appropriately and on time"], "things_needed": []},
        {"name": "Restore Service", "responsible_party": "IT Operations",
         "tasks": ["Restore affected systems/services from clean backups or rebuilds",
                    "Validate systems are clean before reconnecting them to the network"],
         "objectives": ["Normal service is restored safely"], "things_needed": []},
        {"name": "Address Root Cause", "responsible_party": "IT Operations / Security",
         "tasks": ["Fix the underlying vulnerability, misconfiguration, or process gap that allowed this",
                    "Consider technology, process, and personnel factors, not just the technical fix"],
         "objectives": ["The root cause is actually remediated, not just the symptom"], "things_needed": []},
        {"name": "Finalize Incident Analysis Report", "responsible_party": "Incident Commander",
         "tasks": ["Complete a detailed, bias-free report of what happened, the response, and root cause",
                    "Present the report to leadership", "Store the report per your retention/compliance requirements"],
         "objectives": ["A complete incident analysis report is written, presented, and retained"], "things_needed": []},
        {"name": "Post-Mortem", "responsible_party": "Incident Commander + IR staff",
         "tasks": ["Review what worked, what didn't, and what should change",
                    "Assign owners and due dates to any follow-up improvement actions"],
         "objectives": ["Lessons learned are captured and follow-up actions are assigned"], "things_needed": []},
    ]
    return [{"id": _id(), "order": i, **p} for i, p in enumerate(phases)]


def _default_roles_config() -> dict:
    roles = [
        {"name": "Triage Leader", "kind": "standing", "description": "First point of escalation; owns intake and initial classification.",
         "tasks": ["Confirm the initial report and gather basic facts", "Assign initial classification", "Escalate to Incident Commander if Critical/Significant"]},
        {"name": "Incident Commander", "kind": "mandatory", "description": "Owns the incident end-to-end once classified; runs briefings and the Incident Action Plan.",
         "tasks": ["Stand up the response team and assign roles", "Run regular status briefings", "Own the Incident Action Plan", "Sign off before case closure"]},
        {"name": "IT Support/Operations Representative", "kind": "mandatory", "description": "Executes technical containment, evidence collection, and restoration.",
         "tasks": ["Isolate/contain affected systems", "Collect and preserve technical evidence", "Execute restoration plan", "Validate systems before returning to production"]},
        {"name": "Cybersecurity Representative", "kind": "mandatory", "description": "Leads analysis, root cause, and security-specific guidance.",
         "tasks": ["Perform technical analysis of the incident", "Determine root cause", "Recommend containment/eradication steps", "Verify no persistence/backdoors remain"]},
        {"name": "Legal and Compliance Representative", "kind": "mandatory", "description": "Advises on regulatory/legal exposure and required notifications.",
         "tasks": ["Review reporting obligations for applicability", "Advise on legal/regulatory exposure", "Coordinate required external notifications"]},
        {"name": "Public Information Officer", "kind": "optional", "description": "Owns external/internal communications.",
         "tasks": ["Draft internal communications", "Draft external/public communications if needed", "Coordinate messaging with legal"]},
        {"name": "Safety Officer", "kind": "optional", "description": "Ensures responder and physical safety during the incident.",
         "tasks": ["Assess physical safety risks", "Ensure responder safety during containment"]},
        {"name": "Liaison Officer", "kind": "optional", "description": "Coordinates with outside agencies/vendors.",
         "tasks": ["Identify outside agencies/vendors that need coordination", "Maintain point of contact with each"]},
    ]
    return {"roles": [{"id": _id(), **r, "contacts": []} for r in roles], "updated_at": _now_iso()}


# --------------------------------------------------------------------------------
# Wizard: outcome categories + a flat, additive-weights question set. Every question
# is shown once; each option adds weight to one or more categories, optionally adds
# "severity_points" (drives classification independent of category), and can set
# "immediate_containment" (surfaces the fast-path containment phase regardless of
# final category -- mirrors "isolate now" guidance for active ransomware/malware).
# --------------------------------------------------------------------------------
OUTCOME_CATEGORIES = [
    {"id": "phishing", "label": "Phishing / Social Engineering"},
    {"id": "malware", "label": "Malware Infection"},
    {"id": "compromised_credentials", "label": "Compromised Credentials / Account Takeover"},
    {"id": "ransomware", "label": "Ransomware / Active Encryption"},
    {"id": "other_suspicious_activity", "label": "Other Suspicious Activity"},
    {"id": "benign_false_positive", "label": "Likely Benign / False Positive"},
]


def _default_wizard_config() -> dict:
    questions = [
        {
            "text": "What first drew attention to this?",
            "help_text": "Pick whichever best describes how this was first noticed.",
            "options": [
                {"label": "A user reported a suspicious email or link", "weights": {"phishing": 3}},
                {"label": "Antivirus/EDR raised an alert", "weights": {"malware": 3}},
                {"label": "Files or systems appear encrypted, renamed, or show a ransom note", "weights": {"ransomware": 5}, "immediate_containment": True, "severity_points": 5},
                {"label": "Unusual account activity (new-location login, unrequested password reset, MFA prompt)", "weights": {"compromised_credentials": 3}},
                {"label": "Something else, or not sure yet", "weights": {"other_suspicious_activity": 1}},
            ],
        },
        {
            "text": "Did anyone click a link, open an attachment, or enter credentials somewhere unexpected?",
            "options": [
                {"label": "Yes, and they entered credentials on the page", "weights": {"phishing": 2, "compromised_credentials": 3}},
                {"label": "Yes, they clicked/opened it but didn't enter anything", "weights": {"phishing": 3, "malware": 1}},
                {"label": "No", "weights": {}},
                {"label": "Not sure", "weights": {}},
            ],
        },
        {
            "text": "Are any files, folders, or systems currently showing as encrypted, renamed with a strange extension, or displaying a ransom message?",
            "options": [
                {"label": "Yes", "weights": {"ransomware": 5}, "immediate_containment": True, "severity_points": 5},
                {"label": "No", "weights": {}},
                {"label": "Not sure", "weights": {}},
            ],
        },
        {
            "text": "How many people or devices does this appear to affect?",
            "options": [
                {"label": "Just one", "weights": {}, "severity_points": 0},
                {"label": "A few (2-5)", "weights": {}, "severity_points": 2},
                {"label": "Many (6+)", "weights": {}, "severity_points": 4},
                {"label": "Not sure", "weights": {}, "severity_points": 1},
            ],
        },
        {
            "text": "Could sensitive data (customer records, financial data, credentials, health/legal records) have been exposed or accessed?",
            "options": [
                {"label": "Yes, confirmed", "weights": {"phishing": 1, "compromised_credentials": 1}, "severity_points": 5},
                {"label": "Possibly", "weights": {}, "severity_points": 2},
                {"label": "No", "weights": {}, "severity_points": 0},
                {"label": "Not sure", "weights": {}, "severity_points": 1},
            ],
        },
        {
            "text": "Is a critical or public-facing service currently down or behaving abnormally because of this?",
            "options": [
                {"label": "Yes", "weights": {}, "severity_points": 3},
                {"label": "No", "weights": {}, "severity_points": 0},
            ],
        },
        {
            "text": "Did this involve an email or message impersonating a real person, vendor, or executive asking for money, credentials, or sensitive data?",
            "options": [
                {"label": "Yes", "weights": {"phishing": 4}},
                {"label": "No", "weights": {}},
                {"label": "Not sure", "weights": {}},
            ],
        },
        {
            "text": "Have you noticed unusual outbound network traffic, unfamiliar scheduled tasks/processes, or new unrecognized software?",
            "options": [
                {"label": "Yes", "weights": {"malware": 3, "other_suspicious_activity": 1}},
                {"label": "No", "weights": {}},
                {"label": "Not sure", "weights": {}},
            ],
        },
        {
            "text": "Was a password reset, MFA prompt, or account lockout triggered that the user did not initiate?",
            "options": [
                {"label": "Yes", "weights": {"compromised_credentials": 4}},
                {"label": "No", "weights": {}},
                {"label": "Not sure", "weights": {}},
            ],
        },
        {
            "text": "Overall, how confident are you that this is a real security incident rather than routine/expected activity?",
            "options": [
                {"label": "Very confident it's real", "weights": {}, "severity_points": 2},
                {"label": "Somewhat confident", "weights": {}, "severity_points": 1},
                {"label": "Not sure", "weights": {}},
                {"label": "Probably nothing", "weights": {"benign_false_positive": 5}},
            ],
        },
    ]
    for q in questions:
        q["id"] = _id()
        for opt in q["options"]:
            opt["id"] = _id()

    action_plans = {
        "phishing": {
            "summary": "Likely phishing / social engineering.",
            "immediate_actions": [
                "Do not forward the suspicious email further -- report it through your phishing-report process instead",
                "If credentials were entered anywhere, reset that account's password and revoke active sessions/tokens now",
                "Check for suspicious mail rules, forwarding rules, or OAuth app grants added to the mailbox",
                "Warn anyone else who may have received the same message",
            ],
        },
        "malware": {
            "summary": "Likely malware infection.",
            "immediate_actions": [
                "Isolate the affected device from the network (disconnect Wi-Fi/Ethernet) without powering it off",
                "Do not run untrusted 'cleanup' tools before evidence is collected",
                "Capture what the AV/EDR alert actually detected (file path, hash, process) for analysis",
                "Check whether the same indicator has appeared on other devices",
            ],
        },
        "compromised_credentials": {
            "summary": "Likely compromised credentials / account takeover.",
            "immediate_actions": [
                "Reset the account's password and revoke all active sessions/tokens immediately",
                "Check recent sign-in activity, mailbox rules, and OAuth grants for anything unfamiliar",
                "Enable/verify MFA on the account if it wasn't already required",
                "Check whether this account had access to other systems that may also need review",
            ],
        },
        "ransomware": {
            "summary": "Likely ransomware / active encryption in progress.",
            "immediate_actions": [
                "Isolate every affected system from the network immediately -- this is time-critical",
                "Do not pay any ransom or negotiate before Legal/Compliance and the Incident Commander are involved",
                "Do not power off affected systems where memory evidence may still be needed",
                "Identify and disconnect backups from the network if they are reachable from affected systems",
            ],
        },
        "other_suspicious_activity": {
            "summary": "Suspicious activity that doesn't clearly match a specific category yet.",
            "immediate_actions": [
                "Document exactly what was observed, when, and by whom",
                "Avoid taking destructive action until a human responder has reviewed this",
                "Escalate to the security team for manual triage",
            ],
        },
        "benign_false_positive": {
            "summary": "Likely benign or a false positive, but logged for the record.",
            "immediate_actions": [
                "No urgent action needed based on the answers given",
                "If anything changes (new symptoms, more people affected), reopen or re-run the wizard",
            ],
        },
    }
    return {"categories": OUTCOME_CATEGORIES, "questions": questions, "action_plans": action_plans, "updated_at": _now_iso()}


DEFAULT_TOOL_CATALOG = [
    {"name": "Endpoint Detection & Response (EDR)", "applicable_categories": ["malware", "ransomware", "other_suspicious_activity"],
     "description": "Console used to isolate a host from the network and pull process/file details.", "location": "",
     "linked_integration": "CrowdStrike Falcon Spotlight", "vendor": ""},
    {"name": "Vulnerability Management (VMDR)", "applicable_categories": ["malware", "ransomware", "other_suspicious_activity"],
     "description": "Confirms whether the exploited vulnerability is already known/tracked and whether it's patched elsewhere.", "location": "",
     "linked_integration": "Qualys VMDR", "vendor": ""},
    {"name": "Email Security / Phishing Report Button", "applicable_categories": ["phishing"],
     "description": "Where suspicious emails get reported and quarantined org-wide.", "location": "", "linked_integration": None, "vendor": ""},
    {"name": "Identity Provider / SSO Admin Console", "applicable_categories": ["compromised_credentials", "phishing"],
     "description": "Used to reset passwords, revoke sessions, and review sign-in logs.", "location": "",
     "linked_integration": None, "vendor": "Google Workspace Admin Console"},
    {"name": "Backup / Recovery System", "applicable_categories": ["ransomware"],
     "description": "Where clean restore points live -- confirm it's not reachable from affected systems.", "location": "",
     "linked_integration": None, "vendor": ""},
    {"name": "Network Firewall / Segmentation Console", "applicable_categories": ["ransomware", "malware", "other_suspicious_activity"],
     "description": "Used to isolate affected network segments quickly.", "location": "", "linked_integration": None, "vendor": "Palo Alto"},
    {"name": "On-Prem Domain Controllers", "applicable_categories": ["compromised_credentials", "ransomware", "malware"],
     "description": "Check for suspicious account/group changes, unusual authentication events, or new scheduled tasks pushed via GPO.", "location": "",
     "linked_integration": None, "vendor": "Active Directory"},
]


def score_wizard(wizard_config: dict, answers: list) -> dict:
    """answers: [{question_id, option_id}, ...]. Returns the scoring result: per-
    category totals/percentages, the winning category, confidence, severity_points,
    and whether an immediate-containment fast path was triggered by any answer."""
    questions_by_id = {q["id"]: q for q in wizard_config["questions"]}
    category_totals = {c["id"]: 0 for c in wizard_config["categories"]}
    severity_points = 0
    immediate_containment = False
    answered = []

    for a in answers:
        q = questions_by_id.get(a.get("question_id"))
        if not q:
            continue
        opt = next((o for o in q["options"] if o["id"] == a.get("option_id")), None)
        if not opt:
            continue
        for cat_id, w in (opt.get("weights") or {}).items():
            if cat_id in category_totals:
                category_totals[cat_id] += w
        severity_points += opt.get("severity_points", 0)
        if opt.get("immediate_containment"):
            immediate_containment = True
        answered.append({"question": q["text"], "answer": opt["label"]})

    total = sum(max(0, v) for v in category_totals.values())
    percentages = {
        cat_id: round((max(0, v) / total) * 100, 1) if total > 0 else 0.0
        for cat_id, v in category_totals.items()
    }
    if total > 0:
        top_category = max(category_totals, key=lambda k: category_totals[k])
        confidence_pct = percentages[top_category]
    else:
        top_category = "other_suspicious_activity"
        confidence_pct = 0.0

    # Classification from severity_points, with a floor of "Critical" if the
    # immediate-containment fast path was triggered by any answer (active
    # ransomware/malware spread shouldn't ever come out below Critical).
    if severity_points >= 8:
        classification = "Critical"
    elif severity_points >= 5:
        classification = "Significant"
    elif severity_points >= 3:
        classification = "Moderate"
    elif severity_points >= 1:
        classification = "Minor"
    else:
        classification = "Negligible"
    if immediate_containment:
        classification = "Critical"

    return {
        "category_totals": category_totals, "category_percentages": percentages,
        "top_category": top_category, "confidence_pct": confidence_pct,
        "severity_points": severity_points, "classification": classification,
        "immediate_containment": immediate_containment, "answered": answered,
    }


async def push_case_event_to_sheet(db, case: dict, event: dict) -> None:
    """Best-effort, one-way export of a single timeline event to a Google Sheet, via
    a Google Apps Script Web App URL the org deploys themselves (see module
    docstring). Never raises -- a Sheets outage shouldn't block IR work. Falls back
    to the connector's default endpoint if this case doesn't have its own override."""
    import logging
    import httpx
    logger = logging.getLogger("vulnops.ir_sheets")

    webhook_url = case.get("sheets_webhook_url")
    shared_secret = None
    if not webhook_url:
        integration = await db.integrations.find_one({"name": "Google Sheets"}, {"_id": 0})
        cfg = (integration or {}).get("config") or {}
        webhook_url = cfg.get("endpoint")
        shared_secret = cfg.get("api_key")
    if not webhook_url:
        return  # not configured -- silently skip, this is optional

    payload = {
        "secret": shared_secret, "case_id": case["id"], "case_number": case.get("case_number"),
        "case_title": case.get("title"), "classification": case.get("classification"),
        "event_type": event.get("type"), "text": event.get("text"), "author": event.get("author"),
        "timestamp": event.get("created_at"),
    }
    try:
        async with httpx.AsyncClient(timeout=15) as c:
            await c.post(webhook_url, json=payload)
    except Exception as e:
        logger.warning(f"Google Sheets export failed for case {case['id']}: {e}")


def build_closure_report(case: dict, phase_progress: list, events: list, evidence: list) -> dict:
    """Compiles everything captured on a case into a structured closure report --
    paraphrased structure of a typical incident-analysis-report + post-mortem +
    management-briefing combo: what happened, the timeline, the response taken, root
    cause, and follow-up actions. Returned as sections the IR lead can review/edit
    before approving, not a final PDF -- the frontend renders this and lets the lead
    tweak free-text fields before Approve locks it in."""
    timeline_lines = [
        f"{e.get('created_at','')[:19]} — {e.get('author','system')}: {e.get('text','')}"
        for e in sorted(events, key=lambda e: e.get("created_at") or "")
    ]
    phases_completed = [p for p in phase_progress if p.get("completed_at")]
    return {
        "case_id": case["id"], "case_number": case.get("case_number"), "title": case.get("title"),
        "classification": case.get("classification"), "outcome_category": case.get("outcome_category"),
        "opened_at": case.get("opened_at"), "closed_at": case.get("closed_at"),
        "summary": case.get("initial_intake") or "",
        "timeline_text": "\n".join(timeline_lines) or "(no timeline events recorded)",
        "phases_completed": [p.get("phase_name") for p in phases_completed],
        "phases_total": len(phase_progress),
        "evidence_manifest": [{"item_no": ev.get("item_no"), "description": ev.get("description")} for ev in evidence],
        "root_cause": case.get("root_cause") or "",
        "follow_up_actions": case.get("follow_up_actions") or [],
        "generated_at": _now_iso(),
    }


def compute_case_progress(case: dict, phase_progress: list, roles_cfg: list, report: dict = None) -> dict:
    """Rolls up everything trackable on a case (phase checklist tasks, each
    assigned role's own checklist, and open/closed follow-up actions) into one
    0-100 progress number for the ticket header. A case only reads 100% once it's
    closed, the closure report is approved, and there are no open follow-up
    actions left -- matches the "approved but follow-ups still open" case the
    follow-up-actions banner also surfaces, so the two stay consistent."""
    phase_total = sum(len(p.get("tasks") or []) for p in phase_progress)
    phase_done = sum(len(p.get("tasks_done") or []) for p in phase_progress)

    roles_by_id = {r["id"]: r for r in (roles_cfg or [])}
    assigned_roles = case.get("assigned_roles") or {}
    role_total = 0
    role_done = 0
    for role_id, assignment in assigned_roles.items():
        role = roles_by_id.get(role_id)
        if not role or not (assignment or {}).get("name"):
            continue
        n = len(role.get("tasks") or [])
        role_total += n
        role_done += len(set((assignment or {}).get("tasks_done") or []) & set(range(n)))

    follow_ups = case.get("follow_up_actions") or []
    follow_total = len(follow_ups)
    follow_done = sum(1 for a in follow_ups if isinstance(a, dict) and a.get("done"))

    weighted_total = phase_total + role_total + follow_total
    weighted_done = phase_done + role_done + follow_done

    if case.get("status") == "closed" and report and report.get("status") == "approved" and follow_done == follow_total:
        pct = 100.0
    elif weighted_total == 0:
        pct = 100.0 if case.get("status") == "closed" else 0.0
    else:
        pct = round((weighted_done / weighted_total) * 100, 1)

    return {
        "pct": pct,
        "phase_tasks_done": phase_done, "phase_tasks_total": phase_total,
        "role_tasks_done": role_done, "role_tasks_total": role_total,
        "follow_ups_done": follow_done, "follow_ups_total": follow_total,
        "fully_resolved": case.get("status") == "closed" and bool(report) and report.get("status") == "approved" and follow_done == follow_total,
    }


# --------------------------------------------------------------------------------
# Mandatory reporting obligations -- modeled on a real trigger-based routing table
# (which agency/stakeholder gets told what, on what legal/contractual deadline, for
# a given kind of incident) an operator shared as a baseline. Kept fully admin-
# editable (a real CRUD collection, not a fixed list) since every org's actual
# regulatory footprint differs -- what's seeded below is a reasonable starting set
# covering common categories (federal critical-infrastructure/ransomware reporting,
# sector-specific regulators, health/payment/criminal-justice data rules, state
# breach-notification law, and plain internal/stakeholder notification) but every
# entry can be edited, disabled, or removed, and orgs should always confirm current
# requirements with legal/compliance -- these are starting points, not legal advice.
# --------------------------------------------------------------------------------
def _default_reporting_obligations() -> list:
    obligations = [
        {"name": "Ransomware Payment (CISA/CIRCIA)",
         "trigger_description": "Your org executes or facilitates a ransomware payment following an extortion demand.",
         "reporting_target": "Cybersecurity and Infrastructure Security Agency (CISA) via CIRCIA",
         "timeline_hours": 24, "timeline_text": "Within 24 hours of payment", "contacts": []},
        {"name": "Critical Infrastructure Incident (CISA/CIRCIA)",
         "trigger_description": "A substantial cyber incident affecting critical infrastructure you operate (e.g. water/wastewater, emergency services/911, public health, power).",
         "reporting_target": "Cybersecurity and Infrastructure Security Agency (CISA) via CIRCIA",
         "timeline_hours": 72, "timeline_text": "Within 72 hours of reasonable belief an incident occurred", "contacts": []},
        {"name": "Transportation / Airport Security Incident",
         "trigger_description": "A cyber incident impacting airport networks, air traffic control interfaces, security screening infrastructure, or credentialing systems.",
         "reporting_target": "Internal airport security team / TSA / FAA",
         "timeline_hours": 24, "timeline_text": "Immediate notification (typically within 24 hours per applicable security directives)", "contacts": []},
        {"name": "Elections Infrastructure Incident",
         "trigger_description": "A cyber incident or unauthorized access involving voting systems, voter registration databases, or election-night reporting software.",
         "reporting_target": "Secretary of State / State elections office / EI-ISAC",
         "timeline_hours": None, "timeline_text": "Immediate (per state election security protocols)", "contacts": []},
        {"name": "State Government Mandate",
         "trigger_description": "Operational disruption of public services, administrative networks, or government information systems.",
         "reporting_target": "State CIO / State cyber command / State Attorney General",
         "timeline_hours": 72, "timeline_text": "Typically 24-72 hours (varies by state statute)", "contacts": []},
        {"name": "Health Data Breach — 500+ records (HIPAA)",
         "trigger_description": "Breach of unsecured Protected Health Information (PHI) affecting 500 or more individuals.",
         "reporting_target": "HHS Office for Civil Rights (OCR) and prominent media outlets",
         "timeline_hours": 24 * 60, "timeline_text": "Without unreasonable delay, no later than 60 days from discovery", "contacts": []},
        {"name": "Health Data Breach — under 500 records (HIPAA)",
         "trigger_description": "Breach of unsecured PHI affecting fewer than 500 individuals.",
         "reporting_target": "HHS Office for Civil Rights (OCR)",
         "timeline_hours": None, "timeline_text": "No later than 60 days after the end of the calendar year", "contacts": []},
        {"name": "Criminal Justice Information Incident (CJIS)",
         "trigger_description": "A security incident impacting systems housing Criminal Justice Information (CJI).",
         "reporting_target": "State CJIS Systems Officer (CSO) / FBI CJIS",
         "timeline_hours": None, "timeline_text": "Immediate upon identification", "contacts": []},
        {"name": "Payment Card Data Breach (PCI-DSS)",
         "trigger_description": "Breach of a Cardholder Data Environment (CDE) used for payments.",
         "reporting_target": "Merchant acquiring bank / payment card brands",
         "timeline_hours": None, "timeline_text": "Immediate (per merchant agreement)", "contacts": []},
        {"name": "PII Breach (State Breach-Notification Law)",
         "trigger_description": "Unauthorized acquisition or compromise of residents' or employees' personally identifiable information (SSNs, driver's licenses, financial accounts).",
         "reporting_target": "Affected individuals / State Attorney General",
         "timeline_hours": None, "timeline_text": "Varies by state law (typically 30-45 days; some require immediate disclosure)", "contacts": []},
        {"name": "Cyber Insurance Carrier Notification",
         "trigger_description": "Any incident that may give rise to a claim under your cyber insurance policy.",
         "reporting_target": "Cyber insurance carrier / broker",
         "timeline_hours": None, "timeline_text": "Per your policy's notice provisions -- confirm with your broker", "contacts": []},
        {"name": "Internal Leadership & Stakeholder Notification",
         "trigger_description": "Any incident classified Moderate or above -- your own org's leadership and affected-department stakeholders need to know, independent of any external regulator.",
         "reporting_target": "Executive leadership / affected department heads",
         "timeline_hours": None, "timeline_text": "Per your internal escalation policy", "contacts": []},
    ]
    now = _now_iso()
    return [{"id": _id(), **o, "auto_notify": False, "active": True, "created_at": now} for o in obligations]


async def send_obligation_notification(db, case: dict, instance: dict) -> dict:
    """Notifies every contact on an attached obligation via email and/or cell (SMS)
    depending on what's on file for them, plus a webhook if one is set -- routed
    through notifier.deliver() (the same delivery + outbox-logging pipeline the
    Notifications module uses for every other trigger in the app) instead of a
    separate ad hoc email/webhook implementation, so there's one delivery code path,
    one outbox log, and one place SMTP/SMS-gateway config lives. Contacts are a
    custom per-obligation list (not admin-configured channels), so synthetic
    one-off channel dicts are built here and handed to deliver() rather than
    reading from notification_channels. Best-effort per-contact -- one bad
    address/number shouldn't block the rest. Returns a summary the caller logs
    onto the case timeline."""
    from notifier import deliver
    import logging
    logger = logging.getLogger("vulnops.ir_notify")

    ctx = {
        "obligation_name": instance["name"], "case_number": case.get("case_number"), "title": case.get("title"),
        "trigger_description": instance.get("trigger_description", ""), "reporting_target": instance.get("reporting_target", ""),
        "timeline_text": instance.get("timeline_text", ""), "classification": case.get("classification"),
        "summary": (case.get("initial_intake", "") or "")[:500], "url": "",
    }
    sent, failed = [], []
    for contact in (instance.get("contacts") or []):
        for kind, addr in (("email", contact.get("email")), ("sms", contact.get("phone"))):
            if not addr:
                continue
            channel = {"type": kind, "to": addr, "name": contact.get("name") or instance["name"]}
            try:
                result = await deliver(channel, "ir_obligation_notify", ctx, db)
                (sent if result.get("delivered") else failed).append(addr)
            except Exception as e:
                logger.warning(f"Obligation notify ({kind}) failed for {addr}: {e}")
                failed.append(addr)
    webhook_url = instance.get("notify_webhook_url")
    if webhook_url:
        channel = {"type": "webhook", "webhook_url": webhook_url, "name": instance["name"]}
        try:
            result = await deliver(channel, "ir_obligation_notify", ctx, db)
            (sent if result.get("delivered") else failed).append(webhook_url)
        except Exception as e:
            logger.warning(f"Obligation notify webhook failed: {e}")
            failed.append(webhook_url)
    return {"sent": sent, "failed": failed}


def build_case_docx(case: dict, phase_progress: list, events: list, evidence: list,
                     obligations: list, report: dict = None, artifacts: list = None,
                     related_entities: list = None):
    """Builds a real .docx of the case -- opens natively in Word, and Google Docs
    converts .docx on upload/open, so this one export covers both "give it to us as
    a Word doc" and "give it to us as a Google Doc to import" without needing any
    Google OAuth/service-account plumbing. Returns raw bytes (io.BytesIO contents).
    Replaces the earlier Google-Sheets-webhook-as-the-only-option approach, which
    turned out to be more setup than most orgs wanted for what they actually needed:
    something they can just open and read/edit."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title = doc.add_heading(f"{case.get('case_number','')} — {case.get('title','')}", level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Classification: {case.get('classification')}   |   Status: {case.get('status')}   |   "
        f"Opened: {(case.get('opened_at') or '')[:19]}"
        + (f"   |   Closed: {(case.get('closed_at') or '')[:19]}" if case.get("closed_at") else "")
    ).italic = True

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(case.get("initial_intake") or "(no initial intake notes recorded)")

    if case.get("recommended_actions"):
        doc.add_heading("Recommended Actions", level=1)
        for a in case["recommended_actions"]:
            doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Phase Checklist", level=1)
    for p in sorted(phase_progress, key=lambda x: x.get("order", 0)):
        n_tasks = len(p.get("tasks") or [])
        n_done = len(p.get("tasks_done") or [])
        hd = doc.add_heading(f"{p.get('order', 0) + 1}. {p.get('phase_name')} ({n_done}/{n_tasks})", level=2)
        for i, t in enumerate(p.get("tasks") or []):
            mark = "[x]" if i in (p.get("tasks_done") or []) else "[ ]"
            doc.add_paragraph(f"{mark} {t}", style="List Bullet")

    if obligations:
        doc.add_heading("Reporting Obligations", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Obligation", "Reporting Target", "Timeline", "Status"]):
            hdr[i].text = h
        for o in obligations:
            row = table.add_row().cells
            row[0].text = o.get("name", "")
            row[1].text = o.get("reporting_target", "")
            row[2].text = o.get("timeline_text", "")
            row[3].text = o.get("status", "")

    doc.add_heading("Evidence Manifest", level=1)
    if evidence:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["#", "Description", "Location"]):
            hdr[i].text = h
        for ev in evidence:
            row = table.add_row().cells
            row[0].text = str(ev.get("item_no", ""))
            row[1].text = ev.get("description", "")
            row[2].text = ev.get("location", "")
    else:
        doc.add_paragraph("(no evidence logged)")

    doc.add_heading("Timeline & Activity", level=1)
    for e in sorted(events, key=lambda e: e.get("created_at") or ""):
        p = doc.add_paragraph()
        p.add_run(f"{(e.get('created_at') or '')[:19]} — {e.get('author','system')} ").bold = True
        p.add_run(f"({e.get('type','note').replace('_',' ')}): ").italic = True
        p.add_run(e.get("text", ""))
        for att in (e.get("attachments") or []):
            mime = att.get("mime", "")
            data_url = att.get("data_url", "")
            if mime.startswith("image/") and "," in data_url:
                try:
                    import base64
                    raw = base64.b64decode(data_url.split(",", 1)[1])
                    doc.add_picture(io.BytesIO(raw), width=Inches(4))
                except Exception:
                    doc.add_paragraph(f"[screenshot could not be embedded: {att.get('name','?')}]")
            else:
                doc.add_paragraph(f"[attachment: {att.get('name','?')}]")

    if report:
        doc.add_heading("Closure Report", level=1)
        doc.add_paragraph(f"Status: {report.get('status')}")
        doc.add_heading("Root Cause", level=2)
        doc.add_paragraph(case.get("root_cause") or report.get("root_cause") or "(not recorded)")
    follow_ups = case.get("follow_up_actions") or []
    if follow_ups:
        doc.add_heading("Follow-up Actions", level=1)
        for a in follow_ups:
            text = a.get("text", "") if isinstance(a, dict) else str(a)
            done = a.get("done") if isinstance(a, dict) else False
            doc.add_paragraph(f"{'[x]' if done else '[ ]'} {text}", style="List Bullet")

    if related_entities:
        doc.add_heading("Related Assets / Entities", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Type", "Name / Identifier", "Notes"]):
            hdr[i].text = h
        for re_ in related_entities:
            row = table.add_row().cells
            row[0].text = (re_.get("type") or "").replace("_", " ").title()
            row[1].text = re_.get("name", "")
            row[2].text = re_.get("notes", "")

    if artifacts:
        doc.add_heading("Attached Artifacts", level=1)
        doc.add_paragraph(
            "Screenshots embedded above are inline in the Timeline & Activity section. "
            "The files below are attached to the case and available for download from the "
            "case's Artifacts panel (not embedded here since they may not be Word-viewable)."
        )
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["File", "Folder", "Size", "Uploaded by"]):
            hdr[i].text = h
        for a in artifacts:
            row = table.add_row().cells
            row[0].text = a.get("filename", "")
            row[1].text = a.get("folder") or "—"
            size = a.get("size") or 0
            row[2].text = f"{size/1024:.1f} KB" if size < 1024*1024 else f"{size/1024/1024:.1f} MB"
            row[3].text = a.get("uploaded_by", "")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------------
# Case artifacts -- arbitrary files (logs, spreadsheets, PDFs, exports from other
# tools, multi-file folders) attached to a case, stored on disk rather than as
# base64 in Mongo (unlike the small inline screenshot attachments on timeline
# notes) since these can be much larger and more varied in type. Same
# path-containment pattern as backup.py's BACKUP_DIR.
# --------------------------------------------------------------------------------
import os as _os
from pathlib import Path as _Path

ARTIFACTS_DIR = _Path(_os.environ.get("IR_ARTIFACTS_DIR", "/app/ir_artifacts"))
MAX_ARTIFACT_BYTES = 25 * 1024 * 1024  # 25MB/file -- generous for logs/PDFs/spreadsheets, not unlimited


def artifact_disk_path(case_id: str, stored_name: str) -> _Path:
    """Resolves a stored artifact's real path, refusing to resolve outside
    ARTIFACTS_DIR/case_id (stored_name is a generated uuid-based name, never the
    original filename, specifically so this can't be tricked into path traversal)."""
    base = (ARTIFACTS_DIR / case_id).resolve()
    candidate = (base / stored_name).resolve()
    if not str(candidate).startswith(str(base)):
        raise ValueError("Invalid artifact path")
    return candidate
