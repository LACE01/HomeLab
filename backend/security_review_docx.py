"""Security Review report -> editable Word (.docx).

Item 39. Deliberately a DEDICATED generator that maps report sections onto real
Word styles (Heading 1/2, Normal, tables with header shading), NOT an HTML->docx
conversion: HTML conversion produces a document full of literal formatting that
fights the person editing it, while style-mapped output stays editable and
imports cleanly into Google Docs.

Mirrors the on-screen/print report section-for-section so the two never drift:
header block, what was reviewed, risk verdict (inherent -> residual + optional
risk-of-not-adopting), compensating controls (item 27), the 5x5 matrix rendered
as a real Word table (item 23), Recommendation and Decision side by side
(item 24), key findings, conditions, data/systems touched, stakeholder input,
and the technical appendix of questionnaire responses.
"""
import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BAND_COLOR = {"Low": "3B82F6", "Medium": "F59E0B", "High": "F97316", "Critical": "EF4444"}
SEV_COLOR = {"Critical": "EF4444", "High": "F97316", "Medium": "F59E0B", "Low": "3B82F6"}
ANSWER_COLOR = {"no": "EF4444", "partial": "F59E0B", "yes": "16A34A"}


def _html_to_text(html: str) -> str:
    """Rich-text notes are stored as HTML. Word export flattens them to readable
    text rather than embedding markup -- block tags become line breaks, list
    items get a bullet, and everything else is dropped."""
    import html as _html
    import re as _re
    text = html or ""
    text = _re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = _re.sub(r"(?i)</(p|div|pre|h[1-6])>", "\n", text)
    text = _re.sub(r"(?i)<li[^>]*>", "\n• ", text)
    text = _re.sub(r"<[^>]+>", "", text)
    text = _html.unescape(text)
    return _re.sub(r"\n{3,}", "\n\n", text).strip()


def _shade(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _run(paragraph, text, *, bold=False, size=None, color=None, italic=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def _kv_paragraph(doc, label, value):
    p = doc.add_paragraph()
    _run(p, f"{label}: ", bold=True, size=10)
    _run(p, value or "—", size=10)
    return p


def _band_cell(table, row, col, label, band):
    cell = table.cell(row, col)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, label.upper(), size=8, color="666666")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, band or "Not scored", bold=True, size=18, color=BAND_COLOR.get(band, "64748B"))
    _shade(cell, BAND_COLOR.get(band, "E5E7EB") + "" if False else "F8FAFC")


def build_review_docx(data: dict) -> bytes:
    """data == the /report-data payload (review, findings, responses,
    questionnaire, interviews, executive_summary, matrix_points,
    compensating_controls, recommendation)."""
    review = data["review"]
    findings = [f for f in data.get("findings", []) if f.get("status") != "draft"]
    responses = data.get("responses") or []
    questionnaire = data.get("questionnaire") or {}
    interviews = data.get("interviews") or []
    rec = data.get("recommendation") or {}
    decision = review.get("decision") or {}
    conditions = [f for f in findings if f.get("is_condition_of_approval")]

    doc = Document()
    # Letter-ish margins; python-docx defaults to 1" which is what we want.
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)

    # ---------------- header ----------------
    h = doc.add_heading("Security Review Report", level=0)
    for r in h.runs:
        r.font.color.rgb = RGBColor.from_string("111827")
    meta = doc.add_paragraph()
    _run(meta, f"{review.get('review_number', '')}", bold=True, size=10)
    bits = [
        datetime.now().strftime("%Y-%m-%d"),
        f"Reviewer: {review.get('assignee') or '—'}",
    ]
    if review.get("requestor_name"):
        bits.append(f"Requestor: {review['requestor_name']}"
                    + (f" ({review['requestor_department']})" if review.get("requestor_department") else ""))
    _run(meta, "  ·  " + "  ·  ".join(bits), size=9, color="666666")

    p = doc.add_paragraph()
    _run(p, "What was reviewed: ", bold=True, size=10)
    _run(p, f"{review.get('entity_name') or review.get('title')} — {review.get('title')}", size=10)

    # ---------------- PART 1: EXECUTIVE ----------------
    doc.add_heading("Part 1 — Executive summary", level=1)
    p = doc.add_paragraph()
    _run(p, "What was reviewed, what the risk is, and what we recommend. No technical background required.",
         size=9, italic=True, color="666666")

    # ---------------- risk verdict ----------------
    doc.add_heading("Risk verdict", level=2)
    inherent = (review.get("inherent_risk") or {}).get("band")
    residual = (review.get("residual_risk") or {}).get("band")
    not_adopting = (review.get("risk_of_not_adopting") or {}).get("band")

    cols = 3 if not_adopting else 2
    vt = doc.add_table(rows=1, cols=cols)
    vt.alignment = WD_TABLE_ALIGNMENT.CENTER
    vt.style = "Table Grid"
    _band_cell(vt, 0, 0, "Risk if adopted as-is", inherent)
    _band_cell(vt, 0, 1, "Risk with required controls", residual)
    if not_adopting:
        _band_cell(vt, 0, 2, "Risk of NOT adopting", not_adopting)

    if data.get("compensating_controls"):
        doc.add_heading("Compensating controls (what moves inherent → residual)", level=2)
        doc.add_paragraph(data["compensating_controls"], style="Normal")
    if review.get("analyst_override_justification"):
        p = doc.add_paragraph()
        _run(p, "Rating override justification: ", bold=True, size=9)
        _run(p, review["analyst_override_justification"], size=9, italic=True)

    # ---------------- 5x5 matrix ----------------
    points = data.get("matrix_points") or []
    if points:
        doc.add_heading("Risk matrix (likelihood × impact)", level=2)
        mt = doc.add_table(rows=6, cols=6)
        mt.style = "Table Grid"
        hdr = mt.cell(0, 0)
        hdr.text = ""
        _run(hdr.paragraphs[0], "L \\ I", size=8, bold=True)
        for i in range(1, 6):
            c = mt.cell(0, i)
            c.text = ""
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(pp, str(i), size=8, bold=True)
            _shade(c, "F1F5F9")
        for row_idx, likelihood in enumerate([5, 4, 3, 2, 1], start=1):
            c0 = mt.cell(row_idx, 0)
            c0.text = ""
            _run(c0.paragraphs[0], str(likelihood), size=8, bold=True)
            _shade(c0, "F1F5F9")
            for col_idx, impact in enumerate(range(1, 6), start=1):
                cell = mt.cell(row_idx, col_idx)
                cell.text = ""
                score = likelihood * impact
                fill = ("DBEAFE" if score <= 4 else "FEF3C7" if score <= 9
                        else "FFEDD5" if score <= 16 else "FEE2E2")
                _shade(cell, fill)
                here = [pt for pt in points if pt["likelihood"] == likelihood and pt["impact"] == impact]
                if here:
                    pp = cell.paragraphs[0]
                    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run(pp, ", ".join(pt["label"] for pt in here), size=7, bold=True)
        legend = doc.add_paragraph()
        _run(legend, "Impact increases left → right; likelihood increases bottom → top. "
                     + "; ".join(f"{pt['label']}: L{pt['likelihood']}×I{pt['impact']} ({pt['band']})" for pt in points),
             size=8, color="666666")

    # ---------------- confidence ----------------
    qs = data.get("questionnaire_scoring")
    if qs:
        p = doc.add_paragraph()
        _run(p, "Assessment confidence: ", bold=True, size=10)
        bits = [f"{qs['confidence_pct']}%",
                f"based on {qs['applicable_questions']} applicable question(s)"]
        if qs.get("unknown_count"):
            bits.append(f"{qs['unknown_count']} unknown")
        if qs.get("pending_vendor_count"):
            bits.append(f"{qs['pending_vendor_count']} awaiting vendor")
        _run(p, " — ".join([bits[0], ", ".join(bits[1:])]), size=10)
        if qs["confidence_pct"] < 70:
            _run(p, " Treat the ratings above as provisional until the gaps are closed.",
                 size=10, italic=True, color="B45309")

    # ---------------- executive summary ----------------
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(data.get("executive_summary") or "—")

    # ---------------- recommendation vs decision ----------------
    if rec.get("recommendation") or rec.get("why") or rec.get("what_was_reviewed"):
        doc.add_heading("Reviewer recommendation", level=2)
        if rec.get("what_was_reviewed"):
            _kv_paragraph(doc, "What was reviewed", rec["what_was_reviewed"])
        if rec.get("why"):
            _kv_paragraph(doc, "Why", rec["why"])
        p = doc.add_paragraph()
        _run(p, "Recommendation: ", bold=True, size=11)
        _run(p, rec.get("recommendation") or "—", bold=True, size=11, color="2563EB")
        if rec.get("rationale"):
            doc.add_paragraph(rec["rationale"])
        if rec.get("authored_by"):
            p = doc.add_paragraph()
            _run(p, f"— {rec['authored_by']}", size=8, italic=True, color="666666")

    doc.add_heading("Decision", level=2)
    p = doc.add_paragraph()
    _run(p, decision.get("outcome") or "Pending", bold=True, size=12,
         color="16A34A" if decision.get("outcome", "").startswith("Approved") else "111827")
    if decision.get("decision_maker"):
        _kv_paragraph(doc, "Decision maker", decision["decision_maker"])
    if decision.get("decision_date"):
        _kv_paragraph(doc, "Decision date", str(decision["decision_date"])[:10])
    if decision.get("expiration_date"):
        _kv_paragraph(doc, "Approval expires", str(decision["expiration_date"])[:10])
    if decision.get("rationale"):
        doc.add_paragraph(decision["rationale"])
    if rec.get("recommendation") and decision.get("outcome") and \
            rec["recommendation"].strip().lower() not in decision["outcome"].strip().lower():
        p = doc.add_paragraph()
        _run(p, "Note: the decision differs from the reviewer's recommendation.",
             size=9, italic=True, color="B45309")

    # ---------------- conditions ----------------
    if conditions:
        doc.add_heading("Conditions of approval", level=2)
        ct = doc.add_table(rows=1, cols=4)
        ct.style = "Table Grid"
        for i, head in enumerate(["Condition", "Owner", "Deadline", "Status"]):
            c = ct.cell(0, i)
            c.text = ""
            _run(c.paragraphs[0], head, bold=True, size=9)
            _shade(c, "1F2937")
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")
        for cond in conditions:
            row = ct.add_row()
            row.cells[0].text = cond.get("description") or ""
            row.cells[1].text = cond.get("owner") or "—"
            row.cells[2].text = cond.get("condition_deadline") or "—"
            row.cells[3].text = (cond.get("condition_met") or "pending").replace("_", " ")
            for cell in row.cells:
                for pp in cell.paragraphs:
                    for r in pp.runs:
                        r.font.size = Pt(9)

    # ---------------- key findings ----------------
    if findings:
        doc.add_heading("Key findings", level=2)
        for f in findings[:10]:
            p = doc.add_paragraph(style="List Bullet")
            _run(p, f"[{f.get('severity')}] ", bold=True, size=10,
                 color=SEV_COLOR.get(f.get("severity"), "64748B"))
            _run(p, f.get("description") or "", size=10)
            if f.get("recommendation"):
                sub = doc.add_paragraph()
                sub.paragraph_format.left_indent = Inches(0.5)
                _run(sub, f"Recommendation: {f['recommendation']}", size=9, color="555555")

    # ---------------- data & systems touched ----------------
    doc.add_heading("Data & systems touched", level=2)
    classifications = review.get("data_classifications") or []
    _kv_paragraph(doc, "Data classifications", ", ".join(classifications) if classifications else "None selected")
    if review.get("entity_domain"):
        _kv_paragraph(doc, "Vendor domain", review["entity_domain"])
    if review.get("scope_statement"):
        _kv_paragraph(doc, "Scope", review["scope_statement"])

    # ---------------- PART 2: TECHNICAL ----------------
    doc.add_page_break()
    doc.add_heading("Part 2 — Technical detail", level=1)
    p = doc.add_paragraph()
    _run(p, "The evidence behind Part 1: scope, verification results, full questionnaire, working notes, "
            "and supporting documents.", size=9, italic=True, color="666666")

    # ---------------- in-scope assets ----------------
    assets = data.get("linked_assets") or []
    if assets:
        doc.add_heading(f"In-scope assets ({len(assets)})", level=2)
        at = doc.add_table(rows=1, cols=4)
        at.style = "Table Grid"
        for i, head in enumerate(["Host", "Team", "Criticality", "Open findings"]):
            c = at.cell(0, i)
            c.text = ""
            _run(c.paragraphs[0], head, bold=True, size=9)
            _shade(c, "1F2937")
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")
        for a in assets:
            row = at.add_row()
            row.cells[0].text = a.get("hostname") or ""
            row.cells[1].text = a.get("owner_team") or "—"
            row.cells[2].text = a.get("criticality") or "—"
            crit = a.get("critical_high_findings") or 0
            row.cells[3].text = f"{a.get('open_findings', 0)}" + (f" ({crit} crit/high)" if crit else "")
            for cell in row.cells:
                for pp in cell.paragraphs:
                    for r in pp.runs:
                        r.font.size = Pt(9)

    # ---------------- external checks ----------------
    checks = data.get("external_checks") or {}
    if checks.get("company_posture") or checks.get("technical_posture"):
        doc.add_heading("External verification checks", level=2)
        for key, title in (("company_posture", "Company posture"),
                            ("technical_posture", "Technical posture")):
            panel = checks.get(key)
            if not panel:
                continue
            doc.add_heading(title, level=3)
            if panel.get("summary"):
                p = doc.add_paragraph()
                _run(p, panel["summary"].get("headline") or "", size=10)
            for c in panel.get("results", []):
                p = doc.add_paragraph(style="List Bullet")
                _run(p, f"{c.get('label') or c.get('check')}: ", bold=True, size=9.5)
                colour = ("B45309" if c.get("status") == "attention"
                          else "15803D" if c.get("status") == "ok" else "64748B")
                _run(p, f"{c.get('status_plain') or c.get('status')} — ", size=9.5, color=colour)
                _run(p, c.get("summary") or "", size=9.5)
                if c.get("why_it_matters"):
                    sub = doc.add_paragraph()
                    sub.paragraph_format.left_indent = Inches(0.5)
                    _run(sub, f"Why it matters: {c['why_it_matters']}", size=8.5, color="666666")

    # ---------------- stakeholder input ----------------
    if interviews:
        doc.add_heading("Stakeholder input", level=2)
        for it in interviews:
            p = doc.add_paragraph()
            _run(p, f"{it.get('who')} ", bold=True, size=10)
            _run(p, f"({it.get('role') or '—'}, {it.get('when') or '—'}): ", size=9, color="666666")
            _run(p, it.get("summary") or "", size=10)

    # ---------------- notes ----------------
    notes = data.get("notes") or []
    if notes:
        doc.add_heading("Analyst working notes", level=2)
        p = doc.add_paragraph()
        _run(p, "Included because this copy is the internal audit package. Notes are never included in a "
                "shared/external report.", size=8.5, italic=True, color="666666")
        for n in notes:
            p = doc.add_paragraph()
            _run(p, f"{n.get('author')} · {str(n.get('at'))[:19].replace('T', ' ')}", size=8.5, color="666666")
            body = _html_to_text(n.get("html")) if n.get("html") else (n.get("text") or "")
            doc.add_paragraph(body)

    # ---------------- supporting documents ----------------
    attachments = data.get("attachments") or []
    if attachments:
        doc.add_heading(f"Supporting documents ({len(attachments)})", level=2)
        for a in attachments:
            p = doc.add_paragraph(style="List Bullet")
            _run(p, a.get("name") or "", bold=True, size=9.5)
            meta_bits = [a.get("category") or ""]
            if a.get("description"):
                meta_bits.append(a["description"])
            meta_bits.append(f"{round((a.get('size_bytes') or 0) / 1024)} KB")
            meta_bits.append(f"uploaded {str(a.get('uploaded_at'))[:10]} by {a.get('uploaded_by')}")
            _run(p, " — " + ", ".join(x for x in meta_bits if x), size=9, color="666666")

    # ---------------- technical appendix ----------------
    if questionnaire and responses:
        doc.add_heading("Questionnaire responses", level=2)
        by_order = {r["question_order"]: r for r in responses}
        current_domain = None
        for q in questionnaire.get("questions", []):
            r = by_order.get(q["order"])
            if not r:
                continue
            if q.get("domain") != current_domain:
                current_domain = q.get("domain")
                doc.add_heading(current_domain or "General", level=2)
            p = doc.add_paragraph()
            _run(p, f"Q{q['order']}. ", bold=True, size=9)
            _run(p, q.get("text") or "", size=9)
            _run(p, f"  {str(r.get('answer', '')).upper()}", bold=True, size=9,
                 color=ANSWER_COLOR.get(r.get("answer"), "64748B"))
            if r.get("na_reason_code"):
                _run(p, f" ({r['na_reason_code'].replace('_', ' ')})", size=8, color="666666")
            if q.get("cis_mapping"):
                _run(p, f"  [CIS {q['cis_mapping']}]", size=8, color="888888")
            if r.get("auto_answered"):
                _run(p, f"  [{r.get('source_tag', 'auto-answered')}]", size=8, italic=True, color="2563EB")
            if r.get("evidence_text"):
                sub = doc.add_paragraph()
                sub.paragraph_format.left_indent = Inches(0.4)
                _run(sub, r["evidence_text"], size=9, color="555555")

    # ---------------- footer stamp ----------------
    doc.add_paragraph()
    stamp = doc.add_paragraph()
    _run(stamp, f"{review.get('review_number')} · Playbook {review.get('playbook_key')} "
                f"v{review.get('playbook_version')} · Template {review.get('template_key')} "
                f"v{review.get('template_version')} · Generated "
                f"{datetime.now().strftime('%Y-%m-%d %H:%M')}",
         size=8, color="888888")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
